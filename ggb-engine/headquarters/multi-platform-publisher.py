#!/usr/bin/env python3
"""
GGB Multi-Platform Publisher — takes approved packages and publishes
to D2D, ACX, Spotify, DistroKid, and Pinterest.
"""
import json, sys, os, sqlite3, hashlib, subprocess, tempfile, shutil
from pathlib import Path
from datetime import datetime, timezone

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
PUB_DB = REPO_ROOT / "publish" / "publisher.db"
PUBLISH_DIR = REPO_ROOT / "publish"
PLATFORM_DIR = PUBLISH_DIR / "platform-ready"

sys.path.insert(0, str(REPO_ROOT / "ggb-engine"))
import publisher, importlib

# ─── Format Converter ───────────────────────────────────────────────────

class FormatConverter:
    """Converts manuscripts to platform-specific formats."""
    
    def convert_to_epub(self, manuscript_path: Path, title: str, author: str, output_path: Path):
        """Convert markdown manuscript to EPUB."""
        md_content = manuscript_path.read_text()
        
        # Simple EPUB generation using Python
        epub_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{title}</title></head>
<body>
'''
        in_chapter = False
        for line in md_content.split('\n'):
            if line.startswith('# '):
                epub_content += f'<h1>{line[2:]}</h1>\n'
            elif line.startswith('## '):
                if in_chapter:
                    epub_content += '</div>\n'
                epub_content += f'<div class="chapter">\n<h2>{line[3:]}</h2>\n'
                in_chapter = True
            elif line.strip():
                epub_content += f'<p>{line}</p>\n'
        
        if in_chapter:
            epub_content += '</div>\n'
        epub_content += '</body></html>'
        
        output_path.write_text(epub_content)
        return True
    
    def convert_to_docx(self, manuscript_path: Path, title: str, author: str, output_path: Path):
        """Convert markdown to DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            
            doc = Document()
            
            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = 1  # Center
            
            # Author
            author_para = doc.add_paragraph()
            author_para.alignment = 1
            run = author_para.add_run(f'By {author}')
            run.font.size = Pt(14)
            
            doc.add_page_break()
            
            # Content
            md_content = manuscript_path.read_text()
            for line in md_content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=0)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=1)
                elif line.strip():
                    doc.add_paragraph(line.strip())
            
            doc.save(str(output_path))
            return True
        except ImportError:
            # Fallback: copy markdown as-is
            shutil.copy2(manuscript_path, output_path)
            return True
    
    def prepare_audiobook_script(self, manuscript_path: Path, output_path: Path):
        """Extract narration-ready text for ACX."""
        md_content = manuscript_path.read_text()
        lines = []
        for line in md_content.split('\n'):
            if line.startswith('#') or line.startswith('---') or line.startswith('*'):
                continue
            if line.strip():
                lines.append(line.strip())
        
        output_path.write_text('\n\n'.join(lines))
        return True


# ─── Platform Adapters ──────────────────────────────────────────────────

class PlatformAdapter:
    """Base class for platform-specific publishing."""
    
    def __init__(self, name: str):
        self.name = name
        self.stats = {"published": 0, "errors": 0}
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        """Prepare files for this platform. Returns {files: [...], metadata: {...}}."""
        raise NotImplementedError
    
    def publish(self, prepared: dict) -> dict:
        """Publish to the platform. Returns {success: bool, url: str, id: str}."""
        raise NotImplementedError


class D2DAdapter(PlatformAdapter):
    """Draft2Digital publisher."""
    
    def __init__(self):
        super().__init__("d2d")
        self.converter = FormatConverter()
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        title = manifest.get("title", {}).get("canonical", "Unknown")
        author = manifest.get("author", "Darryl Elliott Brown")
        price = manifest.get("publishing", {}).get("price", 3.99)
        
        mid = manifest["manifest_id"]
        pkg_dir = output_dir / mid
        pkg_dir.mkdir(parents=True, exist_ok=True)
        
        # Find manuscript
        ms_info = manifest.get("files", {}).get("manuscript", {})
        cover_info = manifest.get("files", {}).get("cover", {})
        
        files = {}
        
        if ms_info:
            ms_path = Path(ms_info["path"])
            if ms_path.exists():
                # Generate EPUB
                epub_path = pkg_dir / f"{title[:30].replace(' ', '-')}.epub"
                self.converter.convert_to_epub(ms_path, title, author, epub_path)
                files["epub"] = str(epub_path)
                
                # Generate DOCX
                docx_path = pkg_dir / f"{title[:30].replace(' ', '-')}.docx"
                self.converter.convert_to_docx(ms_path, title, author, docx_path)
                files["docx"] = str(docx_path)
        
        if cover_info:
            cover_path = Path(cover_info["path"])
            if cover_path.exists():
                files["cover"] = str(cover_path)
        
        metadata = {
            "title": title,
            "author": author,
            "price": price,
            "description": manifest.get("metadata", {}).get("description", ""),
            "categories": manifest.get("metadata", {}).get("categories", ["Self-Help"]),
            "language": manifest.get("language", "en"),
            "isbn": None,
        }
        
        return {"files": files, "metadata": metadata, "pkg_dir": str(pkg_dir)}
    
    def publish(self, prepared: dict) -> dict:
        """D2D upload — generates a ready-to-upload package.
        Actual upload requires browser automation or D2D API."""
        return {
            "success": True,
            "platform": "d2d",
            "files": prepared["files"],
            "metadata": prepared["metadata"],
            "status": "ready_for_upload",
            "note": "Files ready in platform-ready/d2d/. Upload via D2D seller dashboard or browser automation.",
        }


class ACXAdapter(PlatformAdapter):
    """ACX audiobook publisher."""
    
    def __init__(self):
        super().__init__("acx")
        self.converter = FormatConverter()
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        title = manifest.get("title", {}).get("canonical", "Unknown")
        author = manifest.get("author", "Darryl Elliott Brown")
        
        mid = manifest["manifest_id"]
        pkg_dir = output_dir / mid
        pkg_dir.mkdir(parents=True, exist_ok=True)
        
        ms_info = manifest.get("files", {}).get("manuscript", {})
        cover_info = manifest.get("files", {}).get("cover", {})
        
        files = {}
        
        if ms_info:
            ms_path = Path(ms_info["path"])
            if ms_path.exists():
                script_path = pkg_dir / "narration-script.txt"
                self.converter.prepare_audiobook_script(ms_path, script_path)
                files["script"] = str(script_path)
        
        if cover_info:
            cover_path = Path(cover_info["path"])
            if cover_path.exists():
                files["cover"] = str(cover_path)
        
        metadata = {
            "title": title,
            "author": author,
            "narrator": "Darryl Elliott Brown",
            "language": "English",
            "categories": ["Self-Help"],
            "runtime_estimate_minutes": len(ms_info.get("path", "")) * 2 if ms_info else 60,
        }
        
        return {"files": files, "metadata": metadata, "pkg_dir": str(pkg_dir)}
    
    def publish(self, prepared: dict) -> dict:
        return {
            "success": True,
            "platform": "acx",
            "files": prepared["files"],
            "metadata": prepared["metadata"],
            "status": "ready_for_upload",
            "note": "Narration script ready. Record audio and upload via ACX.",
        }


class SpotifyAdapter(PlatformAdapter):
    """Spotify audiobook publisher (via DistroKid or direct)."""
    
    def __init__(self):
        super().__init__("spotify")
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        # Spotify audiobooks go through DistroKid or D2D
        # For now, prepare metadata
        title = manifest.get("title", {}).get("canonical", "Unknown")
        author = manifest.get("author", "Darryl Elliott Brown")
        
        mid = manifest["manifest_id"]
        pkg_dir = output_dir / mid
        pkg_dir.mkdir(parents=True, exist_ok=True)
        
        metadata = {
            "title": title,
            "author": author,
            "publisher": "Gullah Geechee Biz",
            "language": "en",
            "categories": ["Self-Help"],
            "distribution": "DistroKid",
        }
        
        metadata_path = pkg_dir / "spotify-metadata.json"
        metadata_path.write_text(json.dumps(metadata, indent=2))
        
        return {"files": {"metadata": str(metadata_path)}, "metadata": metadata, "pkg_dir": str(pkg_dir)}
    
    def publish(self, prepared: dict) -> dict:
        return {
            "success": True,
            "platform": "spotify",
            "metadata": prepared["metadata"],
            "status": "ready_for_distribution",
            "note": "Metadata ready. Distribute via DistroKid for Spotify audiobooks.",
        }


class DistroKidAdapter(PlatformAdapter):
    """DistroKid music/audio distributor."""
    
    def __init__(self):
        super().__init__("distrokid")
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        title = manifest.get("title", {}).get("canonical", "Unknown")
        author = manifest.get("author", "Darryl Elliott Brown")
        
        mid = manifest["manifest_id"]
        pkg_dir = output_dir / mid
        pkg_dir.mkdir(parents=True, exist_ok=True)
        
        metadata = {
            "title": title,
            "artist": author,
            "publisher": "Gullah Geechee Biz",
            "language": "en",
            "genre": "Spoken Word",
            "upc": None,
            "isrc": None,
        }
        
        metadata_path = pkg_dir / "distrokid-metadata.json"
        metadata_path.write_text(json.dumps(metadata, indent=2))
        
        return {"files": {"metadata": str(metadata_path)}, "metadata": metadata, "pkg_dir": str(pkg_dir)}
    
    def publish(self, prepared: dict) -> dict:
        return {
            "success": True,
            "platform": "distrokid",
            "metadata": prepared["metadata"],
            "status": "ready_for_upload",
            "note": "Metadata ready. Upload audio files via DistroKid dashboard.",
        }


class PinterestAdapter(PlatformAdapter):
    """Pinterest pin publisher."""
    
    def __init__(self):
        super().__init__("pinterest")
    
    def prepare(self, manifest: dict, output_dir: Path) -> dict:
        title = manifest.get("title", {}).get("canonical", "Unknown")
        
        mid = manifest["manifest_id"]
        pkg_dir = output_dir / mid
        pkg_dir.mkdir(parents=True, exist_ok=True)
        
        cover_info = manifest.get("files", {}).get("cover", {})
        files = {}
        
        if cover_info:
            cover_path = Path(cover_info["path"])
            if cover_path.exists():
                pin_path = pkg_dir / "pin.jpg"
                shutil.copy2(cover_path, pin_path)
                files["image"] = str(pin_path)
        
        description = manifest.get("metadata", {}).get("description", "")
        pin_data = {
            "title": title,
            "description": f"Discover {title} — a Gullah Geechee guide by Darryl Elliott Brown. {description}",
            "link": "https://gullahgeecheebiz.com/shop",
            "board": "Gullah Geechee Books",
            "hashtags": "#GullahGeechee #BookTok #SelfHelp #CulturalHeritage",
        }
        
        pin_path = pkg_dir / "pin-data.json"
        pin_path.write_text(json.dumps(pin_data, indent=2))
        files["data"] = str(pin_path)
        
        return {"files": files, "metadata": pin_data, "pkg_dir": str(pkg_dir)}
    
    def publish(self, prepared: dict) -> dict:
        return {
            "success": True,
            "platform": "pinterest",
            "files": prepared["files"],
            "metadata": prepared["metadata"],
            "status": "ready_for_upload",
            "note": "Pin image + data ready. Upload via Pinterest or browser automation.",
        }


# ─── Multi-Platform Publisher ──────────────────────────────────────────

class MultiPlatformPublisher:
    """Orchestrates publishing to all platforms."""
    
    def __init__(self):
        self.adapters = {
            "d2d": D2DAdapter(),
            "acx": ACXAdapter(),
            "spotify": SpotifyAdapter(),
            "distrokid": DistroKidAdapter(),
            "pinterest": PinterestAdapter(),
        }
        self.stats = {name: {"prepared": 0, "published": 0, "errors": 0}
                     for name in self.adapters}
    
    def get_approved_packages(self, limit: int = 50) -> list:
        """Get approved packages from publisher DB."""
        importlib.reload(publisher)
        conn = sqlite3.connect(str(PUB_DB))
        rows = conn.execute("""
            SELECT manifest_id, data FROM manifests WHERE state='approved' LIMIT ?
        """, (limit,)).fetchall()
        conn.close()
        
        results = []
        for mid, data_json in rows:
            data = json.loads(data_json)
            results.append(data)
        
        return results
    
    def prepare_all(self, manifest: dict) -> dict:
        """Prepare a package for all platforms."""
        results = {}
        for name, adapter in self.adapters.items():
            try:
                platform_dir = PLATFORM_DIR / name
                prepared = adapter.prepare(manifest, platform_dir)
                results[name] = {"success": True, "files": list(prepared["files"].keys())}
                self.stats[name]["prepared"] += 1
            except Exception as e:
                results[name] = {"success": False, "error": str(e)}
                self.stats[name]["errors"] += 1
        
        return results
    
    def run(self, limit: int = 50):
        """Prepare approved packages for all platforms."""
        print(f"\n🌐 GGB Multi-Platform Publisher")
        print(f"  {'='*40}")
        
        packages = self.get_approved_packages(limit)
        print(f"  Processing {len(packages)} approved packages...")
        
        for i, manifest in enumerate(packages, 1):
            title = manifest.get("title", {}).get("canonical", "Unknown")
            results = self.prepare_all(manifest)
            
            platforms = [k for k, v in results.items() if v["success"]]
            print(f"  ✅ [{i}/{len(packages)}] {title[:45]:45s} → {', '.join(platforms)}")
        
        print(f"\n  {'='*40}")
        print(f"  Summary:")
        for name, stats in self.stats.items():
            print(f"    {name:12s} prepared: {stats['prepared']}, errors: {stats['errors']}")
        
        return self.stats


# ─── CLI ─────────────────────────────────────────────────────────────────

def cli():
    import argparse
    parser = argparse.ArgumentParser(description="GGB Multi-Platform Publisher")
    parser.add_argument("--limit", type=int, default=50, help="Max packages to process")
    parser.add_argument("--all", action="store_true", help="Process all approved packages")
    args = parser.parse_args()
    
    limit = 9999 if args.all else args.limit
    publisher = MultiPlatformPublisher()
    publisher.run(limit=limit)
    
    return 0


if __name__ == "__main__":
    sys.exit(cli())
