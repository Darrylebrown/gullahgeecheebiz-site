#!/usr/bin/env python3
"""
Gullah Geechee Biz — Daily Book Generator
One book per day authored by Darryl Elliott Brown, published by Gullah Geechee Biz.
Generates DOCX ready for KDP and Draft2Digital.
"""

import os
import random
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HOME = os.path.expanduser("~")
OUT_DIR = os.path.join(HOME, "ebooks", "daily")
os.makedirs(OUT_DIR, exist_ok=True)

# ─── BRAND CONSTANTS ───
GOLD = RGBColor(0xD4, 0xAF, 0x37)
NAVY = RGBColor(0x0A, 0x14, 0x28)
CREAM = RGBColor(0xF5, 0xF0, 0xE6)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

# ─── BOOK TOPICS (Gullah Geechee Culture) ───
BOOK_TOPICS = [
    {
        "title": "The Gullah Geechee Kitchen: Recipes from the Lowcountry",
        "subtitle": "A Culinary Journey Through the Sea Islands",
        "chapters": [
            "Introduction: The Gullah Geechee Table",
            "The Rice Kingdom: From West Africa to the Lowcountry",
            "Seafood from the Coast: Shrimp, Crab, and Oyster Traditions",
            "One-Pot Wonders: Gumbo, Frogmore Stew, and Red Rice",
            "The Sunday Dinner: Collards, Cornbread, and Sweet Potatoes",
            "Sweets and Preserves: Benne Wafers, Peach Cobbler, and Fig Preserves",
            "The Gullah Geechee Garden: Okra, Field Peas, and Greens",
            "Seasonal Feasts: Oyster Roasts, Fish Fries, and Community Dinners",
            "Heritage Grains: Rice, Grits, and the Stories They Carry",
            "Conclusion: Passing the Spoon"
        ]
    },
    {
        "title": "Voices from the Sea Islands: Gullah Geechee Stories",
        "subtitle": "Folktales, Legends, and Oral Traditions",
        "chapters": [
            "Introduction: The Power of the Spoken Word",
            "Brer Rabbit and the Wisdom of the Trickster",
            "The Hag and the Boo Hag: Spirits of the Lowcountry",
            "Roots and Conjure: Spiritual Traditions of the Gullah People",
            "The Singing of the Spirituals: From Praise Houses to the World",
            "The Language of the Sea Islands: Preserving Gullah",
            "Stories of the Water: Creeks, Tides, and the Marsh",
            "Ancestors and Elders: The Keepers of Memory",
            "The Ring Shout: Dance, Prayer, and Community",
            "Conclusion: The Story Continues"
        ]
    },
    {
        "title": "The Gullah Geechee Corridor: A Journey Through History",
        "subtitle": "From North Carolina to Florida, One Culture, One Story",
        "chapters": [
            "Introduction: The Corridor Defined",
            "The North Carolina Coast: Brunswick and the Borderlands",
            "Georgetown: The Rice Kingdom's Heart",
            "Charleston: The Holy City's Gullah Roots",
            "Beaufort: The Spirit of the Lowcountry",
            "Hilton Head and St. Helena: Islands of Resilience",
            "Savannah: The Hostess City's Gullah Heritage",
            "The Golden Isles: Brunswick, Darien, and Sapelo",
            "Jacksonville and the Florida Reach",
            "Conclusion: The Corridor Lives"
        ]
    },
    {
        "title": "The Art of the Gullah Geechee People",
        "subtitle": "Sweetgrass, Iron, Indigo, and the Creative Spirit",
        "chapters": [
            "Introduction: Art as Survival, Art as Identity",
            "Sweetgrass Baskets: Weaving West Africa into the Lowcountry",
            "Philip Simmons and the Ironwork of Charleston",
            "Indigo: The Blue Gold of the Gullah People",
            "Quilting: Patterns of Memory and Resistance",
            "Gullah Painting: Jonathan Green and the Visual Tradition",
            "Woodcarving and the Maritime Arts",
            "The Gullah Geechee Aesthetic: Design Principles of a Culture",
            "Contemporary Gullah Artists: Carrying the Tradition Forward",
            "Conclusion: The Art of a People"
        ]
    },
    {
        "title": "The Gullah Geechee Language: A Living Treasure",
        "subtitle": "Understanding the Only African-American Creole in the United States",
        "chapters": [
            "Introduction: What Is Gullah?",
            "The African Roots: Languages of the Rice Coast",
            "The Birth of a Creole: How Gullah Was Formed",
            "Gullah Grammar: A Simple Guide",
            "Everyday Gullah: Common Words and Phrases",
            "Gullah in the Praise House: Spiritual Language",
            "Gullah Storytelling: The Language of Brer Rabbit",
            "The Gullah Bible: Scripture in the Mother Tongue",
            "Preserving Gullah: Education, Media, and the Future",
            "Conclusion: Speak the Language of Your Ancestors"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to the Lowcountry",
        "subtitle": "Where to Go, What to See, What to Eat",
        "chapters": [
            "Introduction: Welcome to the Gullah Geechee Corridor",
            "St. Helena Island: Penn Center and the Heart of the Culture",
            "Beaufort: History on Every Corner",
            "Hilton Head: Beyond the Resorts",
            "Charleston: The Holy City's Gullah Soul",
            "Savannah: Squares, Oaks, and Gullah Stories",
            "The Sea Islands: Daufuskie, Sapelo, Edisto, and St. Simons",
            "Georgetown: The River Town's Gullah Heritage",
            "The Best Gullah Restaurants in the Lowcountry",
            "Conclusion: Your Journey Awaits"
        ]
    },
    {
        "title": "The Gullah Geechee Spiritual Tradition",
        "subtitle": "Faith, Praise Houses, and the Ring Shout",
        "chapters": [
            "Introduction: The Spirit of the Sea Islands",
            "The Praise House: The Heartbeat of Gullah Faith",
            "The Ring Shout: The Oldest African-American Worship Tradition",
            "Gullah Spirituals: Songs That Moved the World",
            "Roots and Conjure: Spiritual Practices of the Ancestors",
            "The Gullah Bible: Scripture in the Mother Tongue",
            "Baptisms in the Creek: Water and the Spirit",
            "The Preacher and the Congregation: Call and Response",
            "Gullah Funerals: Homegoing Celebrations",
            "Conclusion: The Spirit Lives On"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Sweetgrass Baskets",
        "subtitle": "History, Craft, and the Art of Weaving",
        "chapters": [
            "Introduction: The Basket That Carried a Culture",
            "From West Africa to the Sea Islands: The Journey of Sweetgrass",
            "The Materials: Sweetgrass, Palmetto, Pine Needles, and Bulrush",
            "The Techniques: Coiling, Stitching, and Shaping",
            "The Masters: Celebrated Gullah Basket Weavers",
            "The Sweetgrass Basket Market: From Charleston to the World",
            "Caring for Your Sweetgrass Basket",
            "Learning to Weave: A Beginner's Guide",
            "The Future of Sweetgrass Basketry",
            "Conclusion: Every Coil Tells a Story"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to the Combahee River Raid",
        "subtitle": "Harriet Tubman's Greatest Military Operation",
        "chapters": [
            "Introduction: The Night Freedom Came",
            "The Combahee River: A Waterway of History",
            "Harriet Tubman: The Moses of Her People",
            "The Plan: Colonel Montgomery and the 2nd South Carolina",
            "The Raid: June 2, 1863",
            "The Aftermath: 750 Souls Set Free",
            "The Gullah Geechee Soldiers of the Combahee",
            "The Legacy: Freedom, Memory, and the Fight Continues",
            "Visiting the Combahee River Today",
            "Conclusion: The River Remembers"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Heirs' Property",
        "subtitle": "Land, Family, and the Fight to Preserve a Legacy",
        "chapters": [
            "Introduction: The Land That Made Us",
            "What Is Heirs' Property? Understanding the Tradition",
            "The History: How Gullah Geechee Families Lost Land",
            "The Legal Challenge: Why Heirs' Property Is Vulnerable",
            "The Heirs' Property Crisis: Land Loss in the Gullah Corridor",
            "Protecting Your Land: Legal Tools and Strategies",
            "The Heirs' Property Retention Act: A New Hope",
            "Community Land Trusts: A Gullah Geechee Solution",
            "Stories of Land Preservation: Families Who Held On",
            "Conclusion: The Land Is Our Legacy"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to the Sea Islands",
        "subtitle": "A Traveler's Guide to the Islands That Preserved a Culture",
        "chapters": [
            "Introduction: The Islands of Resilience",
            "St. Helena Island: The Cultural Heart",
            "Hilton Head Island: History Beyond the Resorts",
            "Edisto Island: Untamed Beauty",
            "Daufuskie Island: The Remote Treasure",
            "Sapelo Island: Hog Hammock and the Last Gullah Community",
            "St. Simons Island: The Golden Isle's Gullah Roots",
            "Jekyll Island: History and Nature",
            "Tybee Island: The Beach Town's Gullah Heritage",
            "Conclusion: The Islands Endure"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Philip Simmons",
        "subtitle": "The Ironwork Master of Charleston",
        "chapters": [
            "Introduction: The Man Who Shaped Charleston",
            "Early Life: Growing Up Gullah Geechee in Charleston",
            "Learning the Craft: Apprenticeship with Peter Simmons",
            "The Art of Wrought Iron: Techniques of a Master",
            "The Gates of Charleston: Philip Simmons's Masterpieces",
            "National Heritage Fellowship: America's Highest Honor",
            "Teaching the Next Generation: Passing the Hammer",
            "The Philip Simmons Legacy: Ironwork in the Modern Era",
            "Visiting Philip Simmons's Charleston",
            "Conclusion: The Iron Will of a People"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Robert Smalls",
        "subtitle": "The Enslaved Man Who Became a Congressman",
        "chapters": [
            "Introduction: The Escape That Changed History",
            "Early Life: Born into Slavery in Beaufort",
            "The Plan: Commandering the CSS Planter",
            "The Escape: May 13, 1862",
            "Freedom and Service: The Planter's Hero",
            "Politics: Robert Smalls in the U.S. Congress",
            "The Legacy: Civil Rights and the Gullah Geechee People",
            "Robert Smalls in Beaufort: Monuments and Memory",
            "The Robert Smalls School: Education and Empowerment",
            "Conclusion: The Captain of His Own Destiny"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Penn Center",
        "subtitle": "The First School for Freed People in the South",
        "chapters": [
            "Introduction: The School That Changed Everything",
            "Founding: 1862 and the Port Royal Experiment",
            "The Early Years: Educating the Freed People",
            "Penn School: From Elementary to Industrial Education",
            "The Civil Rights Era: Penn Center as a Safe Haven",
            "Dr. Martin Luther King Jr. at Penn Center",
            "The Gullah Geechee Cultural Heritage Corridor",
            "Penn Center Today: Preservation and Education",
            "Visiting Penn Center: A Guide",
            "Conclusion: The Legacy of Learning"
        ]
    },
    {
        "title": "The Gullah Geechee Guide to Gullah Cuisine",
        "subtitle": "The Original Farm-to-Table Tradition of America",
        "chapters": [
            "Introduction: The Taste of the Lowcountry",
            "The West African Roots of Gullah Cooking",
            "Rice: The Grain That Built a Culture",
            "Seafood: The Bounty of the Coast",
            "One-Pot Cooking: Gumbo, Jambalaya, and Frogmore Stew",
            "Vegetables: Okra, Collards, Field Peas, and Sweet Potatoes",
            "Cornbread and Rice: The Staples of the Gullah Table",
            "Desserts: Benne Wafers, Sweet Potato Pie, and Peach Cobbler",
            "The Gullah Kitchen: Tools, Techniques, and Traditions",
            "Conclusion: The Table Is Set"
        ]
    },
]

def get_todays_book():
    """Deterministic book selection based on date."""
    rng = random.Random(int(date.today().strftime("%Y%m%d")))
    return rng.choice(BOOK_TOPICS)

def create_book(book_data):
    """Generate a full DOCX book."""
    today = date.today()
    doc = Document()
    
    # Page setup — 6x9 trade paperback
    section = doc.sections[0]
    section.page_width = Cm(15.24)
    section.page_height = Cm(22.86)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Georgia'
        h.font.color.rgb = NAVY
        if level == 1:
            h.font.size = Pt(20)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        else:
            h.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
    
    # ─── TITLE PAGE ───
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book_data["title"].upper())
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.font.name = 'Georgia'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book_data["subtitle"])
    run.italic = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Georgia'
    
    # Gold divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("★" * 20)
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Darryl Elliott Brown")
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = 'Georgia'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gullah Geechee Biz")
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    run.font.name = 'Georgia'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Preserving a Culture. Telling a Story.")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Georgia'
    
    doc.add_page_break()
    
    # ─── COPYRIGHT PAGE ───
    p = doc.add_paragraph()
    run = p.add_run(f"Copyright © {today.year} by Darryl Elliott Brown")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    
    p = doc.add_paragraph()
    run = p.add_run("Published by Gullah Geechee Biz")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    
    p = doc.add_paragraph()
    run = p.add_run("All rights reserved. No part of this publication may be reproduced, distributed, or transmitted in any form or by any means without prior written permission.")
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    
    p = doc.add_paragraph()
    run = p.add_run(f"First Edition: {today.strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    
    p = doc.add_paragraph()
    run = p.add_run("ISBN: 978-0-000-00000-0 (Pending)")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    
    doc.add_page_break()
    
    # ─── TABLE OF CONTENTS ───
    doc.add_heading("Table of Contents", level=1)
    for i, chapter in enumerate(book_data["chapters"], 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Chapter {i}: {chapter}")
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ─── CHAPTERS ───
    for i, chapter_title in enumerate(book_data["chapters"], 1):
        doc.add_heading(f"Chapter {i}", level=1)
        doc.add_heading(chapter_title, level=2)
        
        # Generate chapter content
        content = generate_chapter_content(book_data["title"], chapter_title, i, len(book_data["chapters"]))
        
        for paragraph_text in content:
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text)
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Georgia'
            p.paragraph_format.first_line_indent = Cm(0.5)
        
        doc.add_page_break()
    
    # ─── ABOUT THE AUTHOR ───
    doc.add_heading("About the Author", level=1)
    about = [
        "Darryl Elliott Brown is a Gullah Geechee publisher, author, and cultural preservationist. Born and raised in the Lowcountry, he founded Gullah Geechee Biz to ensure that the stories, traditions, and language of the Gullah Geechee people are preserved for generations to come.",
        "Through a growing library of books, daily content across social media, and a commitment to cultural education, Darryl works to remind the world of the depth, beauty, and resilience of the Gullah Geechee people. Every volume is written with reverence for the ancestors and hope for the future.",
        "Gullah Geechee Biz is dedicated to preserving a culture and telling a story — one book, one pin, one post at a time."
    ]
    for para in about:
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Georgia'
        p.paragraph_format.first_line_indent = Cm(0.5)
    
    # ─── SAVE ───
    slug = book_data["title"].lower().replace(" ", "-").replace(":", "").replace(",", "").replace("'", "")
    filename = f"{today.strftime('%Y-%m-%d')}-{slug}.docx"
    filepath = os.path.join(OUT_DIR, filename)
    doc.save(filepath)
    
    # Verify
    import zipfile
    z = zipfile.ZipFile(filepath)
    names = z.namelist()
    has_content = 'word/document.xml' in names
    size = os.path.getsize(filepath)
    z.close()
    
    return {
        "filepath": filepath,
        "filename": filename,
        "title": book_data["title"],
        "subtitle": book_data["subtitle"],
        "chapters": len(book_data["chapters"]),
        "size_kb": round(size / 1024),
        "valid": has_content
    }

def generate_chapter_content(book_title, chapter_title, chapter_num, total_chapters):
    """Generate placeholder chapter content. In production, this would use the LLM."""
    # This is a structural placeholder — the actual content generation
    # happens in the cron job prompt which has access to the LLM
    paragraphs = [
        f"In the rich tapestry of Gullah Geechee culture, every story, every tradition, and every recipe carries the weight of centuries. The Gullah Geechee people of the Sea Islands have preserved more than just their heritage — they have preserved a way of seeing the world that connects West Africa to the Lowcountry in ways that continue to shape American culture.",
        f"This chapter explores {chapter_title.lower()} — a subject that lies at the heart of what it means to be Gullah Geechee. From the marshlands of the coast to the praise houses of the islands, the traditions of the Gullah people are not artifacts of the past. They are living, breathing practices that continue to evolve while remaining rooted in the wisdom of the ancestors.",
        f"The story of the Gullah Geechee people is the story of America itself — the pain of slavery, the resilience of freedom, the beauty of cultural preservation, and the ongoing struggle for recognition and justice. But it is also a story of joy, of community, of food shared around a table, of songs sung in a praise house, of baskets woven with hands that remember Africa.",
        f"As we delve deeper into this subject, we honor not just the facts and figures, but the lived experience of a people who have given so much to the world. The Gullah Geechee contribution to American culture — from language to cuisine, from music to craftsmanship — is immeasurable. And it is our responsibility to ensure that this legacy is not only remembered but celebrated.",
        f"In the chapters that follow, we will explore the many dimensions of Gullah Geechee life and culture. Each chapter is a thread in a larger tapestry, and together they tell the story of a people whose spirit could not be broken and whose culture could not be erased. This is the Gullah Geechee story — and it deserves to be told with the reverence, accuracy, and love it has always deserved."
    ]
    return paragraphs

def main():
    print("=" * 60)
    print("  GULLAH GEECHEE BIZ — DAILY BOOK GENERATOR")
    print(f"  Date: {date.today().strftime('%B %d, %Y')}")
    print("=" * 60)
    print()
    
    book_data = get_todays_book()
    print(f"  Title: {book_data['title']}")
    print(f"  Subtitle: {book_data['subtitle']}")
    print(f"  Chapters: {len(book_data['chapters'])}")
    print()
    
    result = create_book(book_data)
    
    print(f"  ✅ Book generated successfully")
    print(f"  📁 {result['filepath']}")
    print(f"  📏 {result['size_kb']} KB, {result['chapters']} chapters")
    print(f"  ✅ Valid DOCX: {result['valid']}")
    print()
    print("=" * 60)
    
    # Output JSON for pipeline
    import json
    print(json.dumps(result))

if __name__ == "__main__":
    main()
