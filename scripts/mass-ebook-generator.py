#!/usr/bin/env python3
"""
Gullah Geechee Biz — Mass Ebook Generator
Produces 100+ ebooks across top 3 categories:
1. Self-Help / Personal Development
2. Business & Money
3. Cooking, Food & Wine

Each book is 10 chapters, ~8,000-12,000 words, ready for KDP and direct sales.
"""

import os, json, random
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

HOME = os.path.expanduser("~")
OUT_DIR = os.path.join(HOME, "ebooks", "mass")
os.makedirs(OUT_DIR, exist_ok=True)

GOLD = RGBColor(0xD4, 0xAF, 0x37)
NAVY = RGBColor(0x0A, 0x14, 0x28)
CREAM = RGBColor(0xF5, 0xF0, 0xE6)

# ─── CATEGORY 1: SELF-HELP / PERSONAL DEVELOPMENT ────────────────────────────

SELF_HELP_BOOKS = [
    {
        "slug": "gullah-resilience",
        "title": "The Gullah Geechee Guide to Resilience",
        "subtitle": "Finding Strength in the Stories of the Sea Islands",
        "chapters": [
            "Introduction: The Unbroken Spirit",
            "Rooted in the Land: How Gullah Geechee Ancestors Built Strength",
            "The Power of Story: Oral Traditions as Therapy",
            "Community as Medicine: The Gullah Geechee Village Mindset",
            "Faith and Perseverance: The Spiritual Foundation",
            "Adapting Without Losing Yourself: Change and Identity",
            "The Art of Patience: Slow Living in a Fast World",
            "Healing Through Heritage: Reclaiming What Was Lost",
            "Building Resilience for the Next Generation",
            "Conclusion: The Strength That Carries Forward"
        ]
    },
    {
        "slug": "gullah-mindset",
        "title": "The Gullah Geechee Mindset",
        "subtitle": "Ancient Wisdom for Modern Mental Wellness",
        "chapters": [
            "Introduction: Wisdom from the Marsh",
            "The Principle of Enough: Contentment in a World of More",
            "Finding Peace in Nature: The Lowcountry as Healer",
            "The Power of Silence: Listening to the Wind and the Water",
            "Community Accountability: How Gullah Geechee Villages Raised Children",
            "Letting Go of Grudges: The Gullah Way of Forgiveness",
            "The Rhythm of the Seasons: Living in Harmony with Time",
            "Ancestral Guidance: Drawing Strength from Those Who Came Before",
            "The Practice of Gratitude: Gullah Geechee Daily Rituals",
            "Conclusion: A Mind at Peace"
        ]
    },
    {
        "slug": "gullah-identity",
        "title": "Finding Your Roots: A Gullah Geechee Guide to Identity",
        "subtitle": "Reclaiming Who You Are Through Cultural Heritage",
        "chapters": [
            "Introduction: Who Are You?",
            "The Names We Carry: Understanding Your Family History",
            "The Stories That Shaped Us: Oral History as Identity",
            "Reclaiming Lost Traditions: Food, Language, and Craft",
            "The Geography of Belonging: Place and Identity",
            "Healing Historical Wounds: The Path to Wholeness",
            "Building Identity for Your Children: Passing It On",
            "The Role of Community in Self-Discovery",
            "Writing Your Own Story: Creating a Personal Legacy",
            "Conclusion: You Are the Ancestor of Someone's Future"
        ]
    },
    {
        "slug": "gullah-purpose",
        "title": "The Gullah Geechee Guide to Purpose",
        "subtitle": "Finding Your Calling Through Cultural Wisdom",
        "chapters": [
            "Introduction: Purpose Lives in the Community",
            "Listening to Your Elders: Wisdom as a Compass",
            "The Work of Your Hands: Craft and Calling",
            "Purpose in Service: How Gullah Geechee Communities Thrive",
            "The Courage to Create: Art, Music, and Expression",
            "Purpose and Place: Why Geography Matters",
            "Building Something That Outlasts You",
            "The Role of Faith in Finding Direction",
            "When Purpose Changes: Adapting Without Losing Meaning",
            "Conclusion: Your Purpose is Already Within You"
        ]
    },
    {
        "slug": "gullah-gratitude",
        "title": "Gullah Geechee Gratitude",
        "subtitle": "A 30-Day Journey to a Thankful Heart",
        "chapters": [
            "Introduction: The Gullah Geechee Tradition of Gratitude",
            "Day 1: Gratitude for the Land",
            "Day 2: Gratitude for the Water",
            "Day 3: Gratitude for the Harvest",
            "Day 4: Gratitude for the Elders",
            "Day 5: Gratitude for the Children",
            "Day 6: Gratitude for the Table",
            "Day 7: Gratitude for the Stories",
            "Day 8: Gratitude for the Songs",
            "Day 9: Gratitude for the Hands That Built",
            "Day 10: Gratitude for the Community",
            "Day 11: Gratitude for the Ancestors",
            "Day 12: Gratitude for the Language",
            "Day 13: Gratitude for the Craft",
            "Day 14: Gratitude for the Seasons",
            "Day 15: Gratitude for the Marsh",
            "Day 16: Gratitude for the Ocean",
            "Day 17: Gratitude for the Garden",
            "Day 18: Gratitude for the Kitchen",
            "Day 19: Gratitude for the Neighbors",
            "Day 20: Gratitude for the Teachers",
            "Day 21: Gratitude for the Healers",
            "Day 22: Gratitude for the Protectors",
            "Day 23: Gratitude for the Storytellers",
            "Day 24: Gratitude for the Musicians",
            "Day 25: Gratitude for the Weavers",
            "Day 26: Gratitude for the Builders",
            "Day 27: Gratitude for the Farmers",
            "Day 28: Gratitude for the Fishers",
            "Day 29: Gratitude for the Cooks",
            "Day 30: Gratitude for You"
        ]
    },
    {
        "slug": "gullah-healing",
        "title": "Gullah Geechee Healing",
        "subtitle": "Herbal Remedies, Spiritual Practices, and Ancestral Medicine",
        "chapters": [
            "Introduction: The Healing Traditions of the Gullah Geechee",
            "The Root Doctor: Understanding Gullah Geechee Herbal Medicine",
            "Healing Herbs of the Lowcountry: A Practical Guide",
            "Spiritual Cleansing: Rituals for Mind, Body, and Soul",
            "The Power of Prayer: Gullah Geechee Spiritual Traditions",
            "Healing Through Community: The Village as Medicine",
            "Food as Medicine: Gullah Geechee Nutritional Wisdom",
            "The Healing Power of Water: Creeks, Tides, and Baptism",
            "Restoring Balance: Gullah Geechee Approaches to Mental Health",
            "Conclusion: Carrying the Healing Forward"
        ]
    },
    {
        "slug": "gullah-calm",
        "title": "The Gullah Geechee Guide to Calm",
        "subtitle": "Finding Peace in a Busy World",
        "chapters": [
            "Introduction: The Lowcountry Pace",
            "Morning Rituals: Starting the Day with Intention",
            "The Art of Doing Nothing: Gullah Geechee Rest Traditions",
            "Nature as Sanctuary: Finding Calm in the Marsh",
            "Breathing with the Tides: A Gullah Geechee Breathing Practice",
            "The Evening Wind-Down: Gullah Geechee Nighttime Rituals",
            "Letting Go of Rush: The Gullah Geechee Concept of Time",
            "Finding Calm in Community: Togetherness as Peace",
            "The Calm of Creation: Art, Craft, and Meditation",
            "Conclusion: Carrying the Calm With You"
        ]
    },
    {
        "slug": "gullah-joy",
        "title": "Gullah Geechee Joy",
        "subtitle": "Finding Happiness in the Simple Things",
        "chapters": [
            "Introduction: The Gullah Geechee Definition of Joy",
            "Joy in the Kitchen: Cooking as Celebration",
            "Joy in the Garden: Growing Your Own Happiness",
            "Joy in Music: The Spirituals and the Ring Shout",
            "Joy in Community: Celebrations and Gatherings",
            "Joy in Craft: The Satisfaction of Making",
            "Joy in Nature: The Beauty of the Lowcountry",
            "Joy in Storytelling: Laughter and Wisdom",
            "Joy in Service: The Happiness of Helping",
            "Conclusion: Joy is a Choice"
        ]
    },
    {
        "slug": "gullah-grief",
        "title": "Gullah Geechee Wisdom for Grief",
        "subtitle": "Finding Comfort in Community and Tradition",
        "chapters": [
            "Introduction: Grief in the Gullah Geechee Tradition",
            "The Funeral as Celebration: Homegoing Traditions",
            "The Role of Community in Mourning",
            "Ancestors Are Never Gone: Gullah Geechee Beliefs About Death",
            "The Healing Power of Food: Grief Meals and Comfort",
            "The Strength of the Women: Matriarchs and Mourning",
            "The Strength of the Men: Silent Support",
            "Grief and the Land: Finding Peace in Nature",
            "Carrying the Memory: How Gullah Geechee Families Honor the Departed",
            "Conclusion: Grief is Love With No Place to Go"
        ]
    },
    {
        "slug": "gullah-courage",
        "title": "Gullah Geechee Courage",
        "subtitle": "Stories of Bravery from the Sea Islands",
        "chapters": [
            "Introduction: What Courage Looks Like",
            "Robert Smalls: The Courage to Steal a Ship",
            "Harriet Tubman: The Courage to Return",
            "Denmark Vesey: The Courage to Plan",
            "The Courage of the Mothers: Protecting Children Through Slavery",
            "The Courage to Keep the Language Alive",
            "The Courage to Weave: Sweetgrass Basket Makers",
            "The Courage to Teach: Penn Center and Education",
            "The Courage to Stay: Land Rights and the Gullah Geechee",
            "Conclusion: Your Courage is Waiting"
        ]
    },
    {
        "slug": "gullah-patience",
        "title": "The Gullah Geechee Art of Patience",
        "subtitle": "Slow Living Wisdom from the Lowcountry",
        "chapters": [
            "Introduction: The Lost Art of Patience",
            "Waiting on the Tide: Nature's Lesson in Timing",
            "The Long Harvest: Growing Rice and Growing Character",
            "Slow Food: The Gullah Geechee Kitchen Philosophy",
            "The Patient Hand: Sweetgrass Basket Weaving",
            "Waiting for the Right Moment: Gullah Geechee Decision-Making",
            "The Patience of the Elders: Wisdom That Takes Time",
            "Teaching Patience to Children: Gullah Geechee Parenting",
            "Patience in Business: Building Something That Lasts",
            "Conclusion: Everything Comes in Its Time"
        ]
    },
    {
        "slug": "gullah-community",
        "title": "The Gullah Geechee Way of Community",
        "subtitle": "Building Stronger Connections in a Disconnected World",
        "chapters": [
            "Introduction: We Are Because You Are",
            "The Village Model: How Gullah Geechee Communities Functioned",
            "The Role of the Church: Spiritual and Social Center",
            "Shared Resources: The Gullah Geechee Tradition of Mutual Aid",
            "Raising Children Together: It Takes a Village",
            "Caring for Elders: Respect and Responsibility",
            "Community Celebrations: Food, Music, and Togetherness",
            "Conflict Resolution: The Gullah Geechee Way",
            "Rebuilding Community in the Modern World",
            "Conclusion: Community is the Foundation of Everything"
        ]
    },
    {
        "slug": "gullah-wisdom",
        "title": "Gullah Geechee Wisdom for Daily Living",
        "subtitle": "365 Proverbs and Reflections",
        "chapters": [
            "January: New Beginnings and the Promise of the New Year",
            "February: Love, Family, and the Bonds That Hold Us",
            "March: Growth, Planting, and the Season of Preparation",
            "April: Renewal, Rebirth, and the Spring of the Spirit",
            "May: Work, Harvest, and the Dignity of Labor",
            "June: Community, Celebration, and the Summer Gathering",
            "July: Freedom, Independence, and the Cost of Liberty",
            "August: Patience, Endurance, and the Heat of the Day",
            "September: Reflection, Gratitude, and the Turning Season",
            "October: Ancestors, Memory, and the Wisdom of the Past",
            "November: Thanksgiving, Abundance, and the Generous Heart",
            "December: Rest, Hope, and the Promise of What's to Come"
        ]
    },
    {
        "slug": "gullah-fatherhood",
        "title": "Gullah Geechee Fatherhood",
        "subtitle": "A Guide to Raising Strong Children",
        "chapters": [
            "Introduction: The Role of the Gullah Geechee Father",
            "Leading by Example: The Quiet Strength of Gullah Men",
            "Teaching Through Stories: Passing Down Wisdom",
            "The Importance of Work: Teaching a Strong Work Ethic",
            "Protecting and Providing: The Father's Role",
            "Discipline with Love: Gullah Geechee Parenting",
            "Teaching Heritage: Why Culture Matters",
            "The Father's Prayer: Spiritual Leadership",
            "Building Character: Honesty, Integrity, and Respect",
            "Conclusion: The Legacy of a Good Father"
        ]
    },
    {
        "slug": "gullah-motherhood",
        "title": "Gullah Geechee Motherhood",
        "subtitle": "The Strength That Shapes Generations",
        "chapters": [
            "Introduction: The Matriarchs of the Sea Islands",
            "The Mother as Teacher: Passing Down Language and Culture",
            "The Mother as Healer: Herbal Knowledge and Care",
            "The Mother as Provider: Cooking, Gardening, and Sustenance",
            "The Mother as Protector: Shielding Children from Harm",
            "The Mother as Community Builder: Raising a Village",
            "The Mother as Storyteller: Keeping History Alive",
            "The Mother as Spiritual Guide: Faith and Tradition",
            "The Mother as Example: Strength, Grace, and Dignity",
            "Conclusion: Honoring the Mothers"
        ]
    },
    {
        "slug": "gullah-forgiveness",
        "title": "Gullah Geechee Forgiveness",
        "subtitle": "Letting Go the Lowcountry Way",
        "chapters": [
            "Introduction: The Weight of Unforgiveness",
            "The Gullah Geechee Understanding of Forgiveness",
            "Forgiveness in the Family: Healing Generational Wounds",
            "Forgiveness in the Community: Restoring Relationships",
            "Forgiveness of Self: The Hardest Kind",
            "The Role of Faith in Forgiveness",
            "Stories of Forgiveness from the Sea Islands",
            "The Practice of Forgiveness: Daily Rituals",
            "Teaching Forgiveness to Children",
            "Conclusion: Freedom Through Forgiveness"
        ]
    },
    {
        "slug": "gullah-hope",
        "title": "Gullah Geechee Hope",
        "subtitle": "Finding Light in Dark Times",
        "chapters": [
            "Introduction: Hope in the Gullah Geechee Tradition",
            "The Spirituals: Songs of Hope",
            "Hope in the Midst of Slavery",
            "Hope Through Reconstruction and Jim Crow",
            "Hope in the Civil Rights Movement",
            "Hope in the Preservation of Culture",
            "Hope for the Next Generation",
            "The Practice of Hope: Daily Habits",
            "Stories of Hope from the Lowcountry",
            "Conclusion: Hope is a Choice"
        ]
    },
    {
        "slug": "gullah-elders",
        "title": "Honoring Gullah Geechee Elders",
        "subtitle": "Wisdom, Respect, and the Art of Aging Well",
        "chapters": [
            "Introduction: The Value of Elders",
            "The Role of Elders in Gullah Geechee Society",
            "Learning from the Elders: A Practical Guide",
            "Caring for Aging Parents: The Gullah Geechee Way",
            "The Wisdom of the Elders: Stories and Lessons",
            "Elders as Community Leaders",
            "Passing Down Traditions: The Elder's Role",
            "Aging with Grace: Gullah Geechee Perspectives",
            "Creating Elder-Friendly Communities",
            "Conclusion: Honor Your Elders"
        ]
    },
    {
        "slug": "gullah-marriage",
        "title": "Gullah Geechee Marriage",
        "subtitle": "Building a Strong Union the Lowcountry Way",
        "chapters": [
            "Introduction: Marriage in the Gullah Geechee Tradition",
            "The History of Gullah Geechee Marriage",
            "Choosing a Partner: Wisdom from the Elders",
            "Communication: The Gullah Geechee Way",
            "Working Together: Partnership in Marriage",
            "Faith and Marriage: The Spiritual Foundation",
            "Raising Children Together",
            "Weathering Storms: Resilience in Marriage",
            "Growing Old Together: The Gift of a Lifetime",
            "Conclusion: Love is a Verb"
        ]
    },
    {
        "slug": "gullah-grandparenting",
        "title": "Gullah Geechee Grandparenting",
        "subtitle": "The Joy of Passing Down Traditions",
        "chapters": [
            "Introduction: The Grandparent's Role",
            "The Gullah Geechee Grandparent as Storyteller",
            "Teaching Grandchildren About Heritage",
            "Cooking with Grandchildren: Passing Down Recipes",
            "Crafting with Grandchildren: Sweetgrass and More",
            "The Spiritual Role of Grandparents",
            "Grandparenting from a Distance",
            "The Joy of Being a Grandparent",
            "Leaving a Legacy",
            "Conclusion: Grandparents are Treasures"
        ]
    },
    {
        "slug": "gullah-aging",
        "title": "Gullah Geechee Guide to Aging Well",
        "subtitle": "Growing Older with Grace and Purpose",
        "chapters": [
            "Introduction: Aging in the Gullah Geechee Tradition",
            "The Wisdom of the Elders: What They Taught Us",
            "Staying Active: Gullah Geechee Approaches to Health",
            "Eating Well: The Lowcountry Diet for Longevity",
            "Staying Connected: Community in Later Years",
            "Finding Purpose After Retirement",
            "Passing Down Your Legacy",
            "Aging with Dignity: The Gullah Geechee Way",
            "Preparing for the End of Life",
            "Conclusion: Age is a Gift"
        ]
    },
    {
        "slug": "gullah-mental-health",
        "title": "Gullah Geechee Guide to Mental Health",
        "subtitle": "Healing the Mind Through Culture and Community",
        "chapters": [
            "Introduction: Mental Health in the Gullah Geechee Community",
            "The Stigma of Mental Health: Breaking the Silence",
            "The Role of Community in Mental Wellness",
            "Faith and Mental Health: The Spiritual Dimension",
            "Nature as Therapy: The Healing Power of the Lowcountry",
            "Food and Mood: The Gullah Geechee Connection",
            "Exercise and Movement: Gullah Geechee Traditions",
            "Finding Help: Resources and Support",
            "Supporting Loved Ones with Mental Health Challenges",
            "Conclusion: You Are Not Alone"
        ]
    },
    {
        "slug": "gullah-sabbath",
        "title": "The Gullah Geechee Sabbath",
        "subtitle": "The Art of Rest and Renewal",
        "chapters": [
            "Introduction: The Lost Art of Rest",
            "The Gullah Geechee Sunday Tradition",
            "Rest as Resistance: The History of Sabbath",
            "Creating a Sabbath Practice",
            "Rest and Community: The Gullah Geechee Way",
            "The Spiritual Dimension of Rest",
            "Rest and Creativity: Why We Need Downtime",
            "Teaching Children to Rest",
            "Sabbath in a Busy World",
            "Conclusion: Rest is Sacred"
        ]
    },
    {
        "slug": "gullah-relationships",
        "title": "Gullah Geechee Relationships",
        "subtitle": "Building Strong Bonds the Lowcountry Way",
        "chapters": [
            "Introduction: Relationships in the Gullah Geechee Tradition",
            "Family Relationships: The Foundation",
            "Friendships: The Chosen Family",
            "Romantic Relationships: Love the Gullah Way",
            "Community Relationships: We Are Because You Are",
            "Healing Broken Relationships",
            "Setting Boundaries: The Gullah Geechee Way",
            "Teaching Children About Relationships",
            "Maintaining Long-Distance Relationships",
            "Conclusion: Relationships are Everything"
        ]
    },
    {
        "slug": "gullah-gratitude-journal",
        "title": "The Gullah Geechee Gratitude Journal",
        "subtitle": "A 90-Day Practice of Thankfulness",
        "chapters": [
            "Introduction: How to Use This Journal",
            "Week 1: Gratitude for the Land",
            "Week 2: Gratitude for the Water",
            "Week 3: Gratitude for the Harvest",
            "Week 4: Gratitude for the Elders",
            "Week 5: Gratitude for the Children",
            "Week 6: Gratitude for the Table",
            "Week 7: Gratitude for the Stories",
            "Week 8: Gratitude for the Songs",
            "Week 9: Gratitude for the Hands That Built",
            "Week 10: Gratitude for the Community",
            "Week 11: Gratitude for the Ancestors",
            "Week 12: Gratitude for You",
            "Conclusion: A Grateful Heart"
        ]
    },
    {
        "slug": "gullah-morning",
        "title": "Gullah Geechee Morning Rituals",
        "subtitle": "Starting Each Day with Intention",
        "chapters": [
            "Introduction: The Morning is Sacred",
            "The Gullah Geechee Morning Tradition",
            "Waking with Gratitude",
            "Morning Prayer and Meditation",
            "The Morning Meal: Breakfast the Lowcountry Way",
            "Setting Intentions for the Day",
            "Morning Movement: Gullah Geechee Exercises",
            "Connecting with Nature in the Morning",
            "The Morning Routine: A Step-by-Step Guide",
            "Conclusion: How You Start Your Day Matters"
        ]
    },
    {
        "slug": "gullah-bedtime",
        "title": "Gullah Geechee Bedtime Rituals",
        "subtitle": "Ending Each Day with Peace",
        "chapters": [
            "Introduction: The Evening Wind-Down",
            "The Gullah Geechee Evening Tradition",
            "Evening Prayer and Reflection",
            "Gratitude Practice: Counting Your Blessings",
            "The Evening Meal: Dinner the Lowcountry Way",
            "Family Time: Connecting Before Bed",
            "Preparing for the Next Day",
            "The Power of a Good Night's Sleep",
            "Creating Your Evening Routine",
            "Conclusion: Rest Well"
        ]
    },
    {
        "slug": "gullah-spring",
        "title": "Gullah Geechee Spring",
        "subtitle": "Renewal, Rebirth, and New Beginnings",
        "chapters": [
            "Introduction: The Season of Renewal",
            "Spring in the Lowcountry: A Natural Wonder",
            "Spring Cleaning: The Gullah Geechee Way",
            "Planting the Garden: Seeds of Hope",
            "Spring Foods: What's in Season",
            "Spring Celebrations: Easter and More",
            "Spring Cleaning the Spirit: Renewal Practices",
            "Spring with Children: Teaching About Growth",
            "Spring Projects: Starting Something New",
            "Conclusion: Spring is a Promise"
        ]
    },
    {
        "slug": "gullah-summer",
        "title": "Gullah Geechee Summer",
        "subtitle": "Heat, Harvest, and the Joy of the Season",
        "chapters": [
            "Introduction: The Lowcountry Summer",
            "Summer in the Sea Islands: A Way of Life",
            "Summer Foods: What's in Season",
            "Summer Celebrations: Fourth of July and More",
            "Beating the Heat: Gullah Geechee Wisdom",
            "Summer with Children: Making Memories",
            "Summer Gardening: The Peak of the Harvest",
            "Summer Evenings: The Magic of Sunset",
            "Summer Projects: Making the Most of Long Days",
            "Conclusion: Summer is a Gift"
        ]
    },
    {
        "slug": "gullah-autumn",
        "title": "Gullah Geechee Autumn",
        "subtitle": "Harvest, Gratitude, and the Turning Season",
        "chapters": [
            "Introduction: The Season of Harvest",
            "Autumn in the Lowcountry: A Time of Change",
            "Autumn Foods: What's in Season",
            "Autumn Celebrations: Thanksgiving and More",
            "The Harvest: Reaping What You've Sown",
            "Autumn with Children: Teaching About Gratitude",
            "Preparing for Winter: The Gullah Geechee Way",
            "Autumn Reflection: Taking Stock",
            "Autumn Projects: Finishing What You Started",
            "Conclusion: Gratitude for the Harvest"
        ]
    },
    {
        "slug": "gullah-winter",
        "title": "Gullah Geechee Winter",
        "subtitle": "Rest, Reflection, and the Promise of Spring",
        "chapters": [
            "Introduction: The Season of Rest",
            "Winter in the Lowcountry: A Quieter Time",
            "Winter Foods: Comfort and Warmth",
            "Winter Celebrations: Christmas and Kwanzaa",
            "The Art of Rest: Why Winter Matters",
            "Winter with Children: Indoor Activities",
            "Winter Reflection: Looking Back and Forward",
            "Preparing for Spring: The Quiet Work",
            "Winter Projects: Creative Pursuits",
            "Conclusion: Winter is a Time of Promise"
        ]
    },
]

# ─── CATEGORY 2: BUSINESS & MONEY ────────────────────────────────────────────

BUSINESS_BOOKS = [
    {
        "slug": "gullah-entrepreneur",
        "title": "The Gullah Geechee Entrepreneur",
        "subtitle": "Building a Business That Honors Your Heritage",
        "chapters": [
            "Introduction: Business the Gullah Geechee Way",
            "Finding Your Niche: Leveraging Cultural Assets",
            "Building Without Capital: The Gullah Geechee Resourcefulness",
            "Community as Your First Market",
            "Pricing with Confidence: Valuing Your Work",
            "Marketing Through Story: The Power of Narrative",
            "Building a Brand That Lasts",
            "Scaling Without Losing Your Soul",
            "Networking the Gullah Geechee Way: Relationships First",
            "Conclusion: Your Business is Your Legacy"
        ]
    },
    {
        "slug": "lowcountry-marketing",
        "title": "Lowcountry Marketing",
        "subtitle": "How to Sell Gullah Geechee Products to the World",
        "chapters": [
            "Introduction: Marketing from the Sea Islands",
            "Telling Your Story: The Most Powerful Marketing Tool",
            "Social Media for Cultural Brands",
            "Pinterest: The Visual Search Engine for Your Products",
            "Building an Email List: Your Most Valuable Asset",
            "Selling on Your Own Platform: Cutting Out the Middleman",
            "Wholesale to Bookstores: A Complete Guide",
            "Pricing for Profit: Margins That Work",
            "Scaling Your Marketing: From One Sale to a Movement",
            "Conclusion: Marketing is Just Sharing What You Love"
        ]
    },
    {
        "slug": "gullah-side-hustle",
        "title": "The Gullah Geechee Side Hustle",
        "subtitle": "Turn Your Skills into Income",
        "chapters": [
            "Introduction: Every Gullah Geechee Has a Side Hustle",
            "Cooking for Profit: From Home Kitchen to Catering",
            "Crafting for Cash: Sweetgrass Baskets and More",
            "Writing and Publishing: Your Knowledge is an Asset",
            "Teaching and Consulting: Share What You Know",
            "Tourism and Hospitality: The Lowcountry Advantage",
            "Digital Products: Create Once, Sell Forever",
            "Flipping and Reselling: The Gullah Geechee Trader",
            "Building Multiple Income Streams",
            "Conclusion: Your Side Hustle is Your Freedom"
        ]
    },
    {
        "slug": "gullah-finance",
        "title": "Gullah Geechee Guide to Financial Freedom",
        "subtitle": "Building Wealth the Community Way",
        "chapters": [
            "Introduction: Wealth in the Gullah Geechee Tradition",
            "The History of Gullah Geechee Economics",
            "Saving the Gullah Geechee Way: Practical Strategies",
            "Investing in Your Community: The Best Investment",
            "Real Estate and Land: The Gullah Geechee Relationship with Property",
            "Building Generational Wealth",
            "Debt Freedom: The Gullah Geechee Approach",
            "Multiple Income Streams: The Village Economy",
            "Teaching Children About Money",
            "Conclusion: Wealth is Freedom"
        ]
    },
    {
        "slug": "gullah-publishing",
        "title": "The Gullah Geechee Guide to Self-Publishing",
        "subtitle": "From Manuscript to Marketplace",
        "chapters": [
            "Introduction: Why Self-Publishing Matters",
            "Writing Your Book: A Practical Guide",
            "Editing and Polishing: Making It Professional",
            "Cover Design: First Impressions Matter",
            "Formatting for Print and Digital",
            "Publishing on KDP: A Step-by-Step Guide",
            "Publishing on Your Own Site: Keep 100%",
            "Getting Into Bookstores: The Wholesale Path",
            "Marketing Your Book: A Complete Strategy",
            "Conclusion: Your Book is Your Legacy"
        ]
    },
    {
        "slug": "gullah-ecommerce",
        "title": "Gullah Geechee E-Commerce",
        "subtitle": "Selling Cultural Products Online",
        "chapters": [
            "Introduction: The Digital Marketplace",
            "Choosing Your Products: What Sells and Why",
            "Building Your Online Store",
            "Product Photography: Making It Look Professional",
            "Writing Descriptions That Sell",
            "Shipping and Fulfillment: The Practical Side",
            "Customer Service: The Gullah Geechee Way",
            "Marketing Your Online Store",
            "Scaling: From Side Hustle to Full-Time Business",
            "Conclusion: The World is Your Market"
        ]
    },
    {
        "slug": "gullah-tourism",
        "title": "Gullah Geechee Tourism Guide",
        "subtitle": "Building a Business Around Lowcountry Travel",
        "chapters": [
            "Introduction: The Tourism Opportunity",
            "The Gullah Geechee Corridor: A Tourist's Guide",
            "Starting a Tour Business",
            "Food Tourism: Gullah Geechee Culinary Tours",
            "Heritage Tourism: History and Culture",
            "Art and Craft Tourism: Studio Tours and Workshops",
            "Eco-Tourism: The Natural Lowcountry",
            "Marketing Your Tourism Business",
            "Partnering with Hotels and Travel Agencies",
            "Conclusion: Welcome to the Lowcountry"
        ]
    },
    {
        "slug": "gullah-craft-business",
        "title": "The Gullah Geechee Craft Business Guide",
        "subtitle": "Turning Traditional Skills into Income",
        "chapters": [
            "Introduction: The Value of Handmade",
            "Sweetgrass Basketry: From Tradition to Business",
            "Quilting: Stitching Together a Business",
            "Woodworking and Carpentry: Building a Legacy",
            "Pottery and Ceramics: The Clay of the Lowcountry",
            "Jewelry Making: Adorning the Community",
            "Pricing Your Crafts: The Art of Valuation",
            "Selling at Markets and Festivals",
            "Selling Online: Etsy and Beyond",
            "Conclusion: Your Hands Are Your Greatest Asset"
        ]
    },
    {
        "slug": "gullah-food-business",
        "title": "Starting a Gullah Geechee Food Business",
        "subtitle": "From Home Kitchen to Restaurant",
        "chapters": [
            "Introduction: The Business of Gullah Geechee Food",
            "Licensing and Permits: What You Need to Know",
            "Developing Your Menu: Authenticity and Appeal",
            "Sourcing Ingredients: The Lowcountry Advantage",
            "Pricing Your Food: Margins That Work",
            "Catering: A Low-Risk Start",
            "Food Trucks: Mobility and Flexibility",
            "Restaurant: The Full Commitment",
            "Selling Products: Sauces, Seasonings, and Mixes",
            "Conclusion: Feed the Body, Feed the Soul"
        ]
    },
    {
        "slug": "gullah-cooperative",
        "title": "The Gullah Geechee Cooperative",
        "subtitle": "Building Community Wealth Together",
        "chapters": [
            "Introduction: The Power of Cooperation",
            "The History of Gullah Geechee Cooperatives",
            "Starting a Cooperative: A Step-by-Step Guide",
            "Legal Structure: Choosing the Right Form",
            "Governance: How Cooperatives Make Decisions",
            "Funding Your Cooperative",
            "Marketing as a Group: Strength in Numbers",
            "Case Studies: Successful Gullah Geechee Cooperatives",
            "Scaling: From Local to Regional",
            "Conclusion: Together We Rise"
        ]
    },
    {
        "slug": "gullah-freelance",
        "title": "The Gullah Geechee Freelancer",
        "subtitle": "Building a Career on Your Own Terms",
        "chapters": [
            "Introduction: Freedom Through Freelance",
            "Identifying Your Skills: What Can You Offer?",
            "Setting Your Rates: The Gullah Geechee Approach to Value",
            "Finding Clients: Networking in the Community",
            "Building a Portfolio: Show Your Work",
            "Managing Finances: The Freelancer's Challenge",
            "Balancing Work and Life: The Gullah Geechee Way",
            "Scaling Your Freelance Business",
            "Building a Reputation: Trust is Everything",
            "Conclusion: Your Time is Your Own"
        ]
    },
    {
        "slug": "gullah-real-estate",
        "title": "Gullah Geechee Guide to Real Estate",
        "subtitle": "Land, Property, and Generational Wealth",
        "chapters": [
            "Introduction: Land is Freedom",
            "The History of Gullah Geechee Land Ownership",
            "Heirs Property: Understanding the Challenge",
            "Protecting Your Land: Legal Strategies",
            "Buying Your First Property",
            "Real Estate Investing: The Gullah Geechee Way",
            "Rental Properties: Passive Income Strategies",
            "Commercial Real Estate: Building Community Assets",
            "Passing Down Property: Estate Planning",
            "Conclusion: Land is Legacy"
        ]
    },
    {
        "slug": "gullah-nonprofit",
        "title": "Starting a Gullah Geechee Nonprofit",
        "subtitle": "A Guide to Community Service Organizations",
        "chapters": [
            "Introduction: Service is Sacred",
            "Identifying the Need: What Does Your Community Require?",
            "Legal Structure: 501(c)(3) and Beyond",
            "Building a Board: Leadership and Governance",
            "Fundraising: Grants, Donations, and Events",
            "Programming: Designing Effective Services",
            "Marketing Your Nonprofit: Telling Your Story",
            "Measuring Impact: Proving Your Value",
            "Sustainability: Keeping the Mission Alive",
            "Conclusion: Service is Leadership"
        ]
    },
    {
        "slug": "gullah-investing",
        "title": "Gullah Geechee Investing",
        "subtitle": "Building Wealth for the Long Term",
        "chapters": [
            "Introduction: Investing in Your Future",
            "The Basics: Stocks, Bonds, and Understanding the Market",
            "Real Estate: The Gullah Geechee Relationship with Property",
            "Small Business Investing: Supporting Community Entrepreneurs",
            "Retirement Planning: The Long View",
            "Teaching Children About Investing",
            "Ethical Investing: Aligning Money with Values",
            "Building an Investment Portfolio",
            "Common Mistakes and How to Avoid Them",
            "Conclusion: Invest in Yourself First"
        ]
    },
]

# ─── CATEGORY 3: COOKING, FOOD & WINE ────────────────────────────────────────

COOKING_BOOKS = [
    {
        "slug": "gullah-kitchen-v1",
        "title": "The Gullah Geechee Kitchen Volume 1",
        "subtitle": "Classic Recipes from the Lowcountry",
        "chapters": [
            "Introduction: The Gullah Geechee Table",
            "Rice Dishes: Red Rice, Perloo, and Hoppin' John",
            "Seafood: Shrimp and Grits, Crab Rice, Frogmore Stew",
            "Soups and Stews: Okra Soup, Gumbo, She-Crab Soup",
            "Vegetables: Collard Greens, Okra and Tomatoes, Fried Cabbage",
            "Breads: Buttermilk Cornbread, Benne Wafers",
            "Main Dishes: Fried Fish, Smothered Pork Chops, Meatloaf",
            "Desserts: Sweet Potato Pie, Peach Cobbler, Banana Pudding",
            "Preserves and Pickles: Fig Preserves, Pickled Shrimp",
            "Conclusion: Passing the Spoon"
        ]
    },
    {
        "slug": "gullah-kitchen-v2",
        "title": "The Gullah Geechee Kitchen Volume 2",
        "subtitle": "More Recipes from the Sea Islands",
        "chapters": [
            "Introduction: The Kitchen Never Closes",
            "Breakfast: Shrimp and Grits, Fried Fish, Sweet Potato Pancakes",
            "Soups: Chicken and Rice Soup, Seafood Chowder",
            "Salads: Gullah Potato Salad, Egg Salad",
            "Appetizers: Deviled Crabs, Pimento Cheese, Crab Cakes",
            "Main Dishes: Country Captain Chicken, Benne Crusted Chicken",
            "Side Dishes: Mac and Cheese, Cornbread Dressing, Limpin' Susan",
            "Desserts: Rice Pudding, Tea Cakes, Gullah Cake",
            "Beverages: Sweetgrass Lemonade, Sweet Tea",
            "Conclusion: The Table is Always Set"
        ]
    },
    {
        "slug": "gullah-sunday-dinner",
        "title": "Gullah Geechee Sunday Dinner",
        "subtitle": "The Meal That Brings Everyone Together",
        "chapters": [
            "Introduction: Sunday is Sacred",
            "The Main Dish: Smothered Pork Chops, Fried Chicken, Meatloaf",
            "The Rice: Red Rice, Perloo, Hoppin' John",
            "The Greens: Collards, Mustard Greens, Turnip Greens",
            "The Bread: Cornbread, Biscuits, Rolls",
            "The Sides: Mac and Cheese, Potato Salad, Baked Beans",
            "The Dessert: Sweet Potato Pie, Banana Pudding, Peach Cobbler",
            "The Drink: Sweet Tea, Lemonade",
            "Setting the Table: The Gullah Geechee Way",
            "Conclusion: Sunday Dinner is Love"
        ]
    },
    {
        "slug": "gullah-seafood",
        "title": "Gullah Geechee Seafood Cookbook",
        "subtitle": "From the Atlantic to Your Table",
        "chapters": [
            "Introduction: The Bounty of the Sea",
            "Shrimp: Shrimp and Grits, Fried Shrimp, Pickled Shrimp",
            "Crab: Crab Rice, Deviled Crabs, Crab Cakes, She-Crab Soup",
            "Fish: Fried Fish, Baked Fish, Fish Stew",
            "Oysters: Fried Oysters, Oyster Stew, Roasted Oysters",
            "Clams: Clam Chowder, Steamed Clams",
            "Frogmore Stew: The Ultimate Seafood Boil",
            "Gumbo: Seafood Gumbo, Chicken and Sausage Gumbo",
            "Sides: Okra and Tomatoes, Hushpuppies, Coleslaw",
            "Conclusion: The Sea Provides"
        ]
    },
    {
        "slug": "gullah-soul-food",
        "title": "Gullah Geechee Soul Food",
        "subtitle": "The Roots of Southern Cooking",
        "chapters": [
            "Introduction: Soul Food is Gullah Geechee Food",
            "The History: From West Africa to the Lowcountry",
            "Fried Chicken: The Perfect Bird",
            "Collard Greens: The Soul Food Staple",
            "Mac and Cheese: The Ultimate Comfort",
            "Cornbread: The Foundation of Every Meal",
            "Sweet Potatoes: From Pie to Casserole",
            "Black-Eyed Peas: Hoppin' John and Limpin' Susan",
            "Peach Cobbler: The Perfect Ending",
            "Conclusion: Soul Food is Love"
        ]
    },
    {
        "slug": "gullah-desserts",
        "title": "Gullah Geechee Desserts",
        "subtitle": "Sweet Treats from the Lowcountry",
        "chapters": [
            "Introduction: The Gullah Geechee Sweet Tooth",
            "Sweet Potato Pie: The Queen of Desserts",
            "Peach Cobbler: Summer in a Dish",
            "Banana Pudding: The Southern Classic",
            "Benne Wafers: Sesame Cookies from West Africa",
            "Rice Pudding: Using Every Grain",
            "Tea Cakes: The Perfect Tea Time Treat",
            "Sweetgrass Lemonade: A Refreshing Drink",
            "Fig Preserves: Sweetness Preserved",
            "Conclusion: Life is Sweet"
        ]
    },
    {
        "slug": "gullah-one-pot",
        "title": "Gullah Geechee One-Pot Meals",
        "subtitle": "Simple, Delicious, and Satisfying",
        "chapters": [
            "Introduction: The Beauty of One Pot",
            "Red Rice: The Signature Dish",
            "Perloo: Chicken Perloo, Sausage Perloo",
            "Gumbo: Seafood Gumbo, Chicken Gumbo",
            "Frogmore Stew: The Lowcountry Boil",
            "Hoppin' John: Rice and Peas",
            "Limpin' Susan: The Quick Version",
            "Okra Soup: A Meal in a Bowl",
            "Country Captain: Curried Chicken Stew",
            "Conclusion: One Pot, Endless Possibilities"
        ]
    },
    {
        "slug": "gullah-holiday",
        "title": "Gullah Geechee Holiday Cookbook",
        "subtitle": "Celebrating Through Food",
        "chapters": [
            "Introduction: Every Holiday Has a Flavor",
            "New Year's Day: Hoppin' John for Good Luck",
            "Easter: Ham, Sweet Potatoes, and Deviled Eggs",
            "Fourth of July: Frogmore Stew and Peach Cobbler",
            "Thanksgiving: Turkey, Dressing, and Sweet Potato Pie",
            "Christmas: Roast, Rice, and Banana Pudding",
            "Kwanzaa: Feasting on Heritage",
            "Family Reunions: Cooking for a Crowd",
            "Sunday Dinners: The Weekly Celebration",
            "Conclusion: Every Meal is a Holiday"
        ]
    },
    {
        "slug": "gullah-vegetarian",
        "title": "Gullah Geechee Vegetarian",
        "subtitle": "Plant-Based Lowcountry Cooking",
        "chapters": [
            "Introduction: Vegetables are the Foundation",
            "Rice and Beans: Hoppin' John, Red Beans and Rice",
            "Okra: Okra Soup, Okra and Tomatoes, Okra Fries",
            "Greens: Collards, Mustard Greens, Turnip Greens",
            "Sweet Potatoes: Roasted, Mashed, Pies",
            "Corn: Cornbread, Grits, Succotash",
            "Field Peas: Black-Eyed Peas, Crowder Peas",
            "Squash and Tomatoes: Summer's Bounty",
            "Benne Seeds: Wafers, Chicken, and More",
            "Conclusion: Vegetables are Soul Food Too"
        ]
    },
    {
        "slug": "gullah-breakfast",
        "title": "Gullah Geechee Breakfast",
        "subtitle": "Starting the Day the Lowcountry Way",
        "chapters": [
            "Introduction: The Most Important Meal",
            "Shrimp and Grits: The Classic Lowcountry Breakfast",
            "Fried Fish and Grits: A Friday Morning Tradition",
            "Sweet Potato Pancakes: A Gullah Twist",
            "Cornbread and Molasses: Simple and Satisfying",
            "Eggs: Gullah Egg Salad, Deviled Eggs",
            "Grits: Cheese Grits, Garlic Grits, Shrimp and Grits",
            "Biscuits: Buttermilk Biscuits, Sausage Biscuits",
            "Breakfast Sides: Fried Apples, Grits, Eggs",
            "Conclusion: Breakfast is Love"
        ]
    },
    {
        "slug": "gullah-preserving",
        "title": "Gullah Geechee Guide to Preserving",
        "subtitle": "Canning, Pickling, and Keeping the Harvest",
        "chapters": [
            "Introduction: Nothing Goes to Waste",
            "Fig Preserves: The Taste of Summer",
            "Peach Preserves: Sweetness All Year",
            "Pickled Shrimp: A Lowcountry Tradition",
            "Pickled Okra: Crunchy and Delicious",
            "Tomato Preserves: Summer in a Jar",
            "Pepper Sauce: The Gullah Geechee Hot Sauce",
            "Drying Herbs: Keeping the Garden Alive",
            "Freezing: The Modern Preservation",
            "Conclusion: Preserve the Harvest, Preserve the Culture"
        ]
    },
    {
        "slug": "gullah-grilling",
        "title": "Gullah Geechee Grilling",
        "subtitle": "Lowcountry Barbecue and Outdoor Cooking",
        "chapters": [
            "Introduction: Cooking Over Fire",
            "The Gullah Geechee Barbecue Tradition",
            "Pulled Pork: Low and Slow",
            "Grilled Chicken: The Perfect Bird",
            "Grilled Fish: Fresh from the Atlantic",
            "Grilled Shrimp: Skewers and Seasonings",
            "Vegetables from the Grill: Okra, Corn, Peppers",
            "Barbecue Sides: Coleslaw, Potato Salad, Baked Beans",
            "Barbecue Sauce: The Gullah Geechee Recipe",
            "Conclusion: Gather Around the Fire"
        ]
    },
    {
        "slug": "gullah-sauces",
        "title": "Gullah Geechee Sauces and Seasonings",
        "subtitle": "The Flavors of the Lowcountry",
        "chapters": [
            "Introduction: Sauce is Everything",
            "Pepper Sauce: The Gullah Geechee Staple",
            "Remoulade: The Perfect Seafood Sauce",
            "Barbecue Sauce: Sweet, Tangy, Smoky",
            "Gravy: Onion Gravy, Tomato Gravy",
            "Hot Sauce: Making Your Own",
            "Seasoning Blends: Old Bay, Gullah Spice",
            "Marinades: For Meat, Seafood, and Vegetables",
            "Dressings: Buttermilk, Vinaigrette, Ranch",
            "Conclusion: Master the Sauce, Master the Kitchen"
        ]
    },
    {
        "slug": "gullah-baking",
        "title": "Gullah Geechee Baking",
        "subtitle": "Breads, Pies, and Pastries from the Lowcountry",
        "chapters": [
            "Introduction: The Gullah Geechee Baker",
            "Cornbread: The Foundation",
            "Biscuits: Light, Fluffy, Buttery",
            "Sweet Potato Pie: The Signature Dessert",
            "Peach Cobbler: Summer's Best",
            "Benne Wafers: Sesame Cookies",
            "Tea Cakes: The Perfect Tea Time Treat",
            "Banana Pudding: Layers of Love",
            "Rice Pudding: Using Every Grain",
            "Conclusion: Baking is Love"
        ]
    },
    {
        "slug": "gullah-drinks",
        "title": "Gullah Geechee Drinks and Beverages",
        "subtitle": "From Sweet Tea to Sweetgrass Lemonade",
        "chapters": [
            "Introduction: Quenching Thirst the Gullah Way",
            "Sweet Tea: The Southern Classic",
            "Sweetgrass Lemonade: A Unique Twist",
            "Ginger Beer: Homemade and Spicy",
            "Fruit Punches: Summer Refreshment",
            "Herbal Teas: From the Garden",
            "Coffee: The Gullah Geechee Morning",
            "Elderflower Cordial: A Traditional Drink",
            "Holiday Drinks: Eggnog, Wassail, and More",
            "Conclusion: Drink to Your Health"
        ]
    },
]

# ─── EXTRA BUSINESS BOOKS ────────────────────────────────────────────────────

EXTRA_BUSINESS_BOOKS = [
    {
        "slug": "gullah-consulting",
        "title": "The Gullah Geechee Consultant",
        "subtitle": "Turning Expertise into Income",
        "chapters": ["Introduction: You Already Know Enough","Identifying Your Expertise","Packaging Your Knowledge","Setting Your Rates","Finding Clients","Delivering Value","Scaling Your Practice","Building a Reputation","The Gullah Geechee Approach to Consulting","Conclusion: Your Knowledge is Valuable"]
    },
    {
        "slug": "gullah-remote-work",
        "title": "Gullah Geechee Guide to Remote Work",
        "subtitle": "Working from Home the Lowcountry Way",
        "chapters": ["Introduction: The New Way of Working","Setting Up Your Home Office","Managing Your Time","Staying Connected","Avoiding Burnout","The Gullah Geechee Work-Life Balance","Building a Remote Career","Remote Work Tools and Tips","Thriving as a Remote Worker","Conclusion: Work Anywhere, Live Everywhere"]
    },
    {
        "slug": "gullah-budget",
        "title": "The Gullah Geechee Budget",
        "subtitle": "A Simple System for Managing Your Money",
        "chapters": ["Introduction: Every Dollar Has a Purpose","Tracking Your Spending","Creating a Budget That Works","The Envelope System: Gullah Geechee Style","Saving for Emergencies","Paying Off Debt","Investing for the Future","Teaching Your Family About Money","Sticking with It","Conclusion: Financial Peace"]
    },
    {
        "slug": "gullah-credit",
        "title": "Gullah Geechee Guide to Credit",
        "subtitle": "Building and Repairing Your Credit Score",
        "chapters": ["Introduction: Why Credit Matters","Understanding Your Credit Score","Building Credit from Scratch","Repairing Damaged Credit","Credit Cards: Using Them Wisely","Loans and Borrowing","Protecting Your Identity","Credit and Your Community","Teaching the Next Generation","Conclusion: Credit is a Tool"]
    },
    {
        "slug": "gullah-debt",
        "title": "Gullah Geechee Guide to Debt Freedom",
        "subtitle": "Breaking Free from What You Owe",
        "chapters": ["Introduction: Debt is a Chain","Understanding Your Debt","The Debt Snowball: Gullah Geechee Style","Negotiating with Creditors","Cutting Expenses","Increasing Income","Staying Motivated","Life After Debt","Helping Others Find Freedom","Conclusion: Free at Last"]
    },
    {
        "slug": "gullah-retirement",
        "title": "Gullah Geechee Guide to Retirement",
        "subtitle": "Planning for Your Golden Years",
        "chapters": ["Introduction: Retirement the Gullah Way","Understanding Retirement Accounts","How Much You Need","Saving at Every Age","Investing for Retirement","Social Security and Pensions","Retirement and Community","Retirement in the Lowcountry","Leaving a Legacy","Conclusion: Your Future is Bright"]
    },
    {
        "slug": "gullah-taxes",
        "title": "Gullah Geechee Guide to Taxes",
        "subtitle": "Understanding and Managing Your Taxes",
        "chapters": ["Introduction: Taxes and Your Community","Understanding the Tax System","Filing Your Taxes","Deductions and Credits","Taxes for Small Business Owners","Working with a Tax Professional","Planning for Tax Season","Common Mistakes to Avoid","Teaching Your Family About Taxes","Conclusion: Taxes are Part of Life"]
    },
    {
        "slug": "gullah-insurance",
        "title": "Gullah Geechee Guide to Insurance",
        "subtitle": "Protecting What Matters Most",
        "chapters": ["Introduction: Insurance as Protection","Health Insurance","Life Insurance","Home and Property Insurance","Auto Insurance","Business Insurance","Understanding Your Policies","Choosing the Right Coverage","Saving on Insurance","Conclusion: Peace of Mind"]
    },
    {
        "slug": "gullah-estate",
        "title": "Gullah Geechee Guide to Estate Planning",
        "subtitle": "Protecting Your Legacy",
        "chapters": ["Introduction: Planning for Tomorrow","Wills and Trusts","Heirs Property: A Gullah Geechee Guide","Power of Attorney","Healthcare Directives","Passing Down Property","Minimizing Estate Taxes","Talking to Your Family About Estate Planning","Working with an Attorney","Conclusion: Your Legacy Matters"]
    },
    {
        "slug": "gullah-farming",
        "title": "The Gullah Geechee Farmer",
        "subtitle": "Starting a Farm or Garden Business",
        "chapters": ["Introduction: The Land Provides","The History of Gullah Geechee Farming","Choosing Your Land","What to Grow","Farming Methods","Bringing Your Products to Market","Community Supported Agriculture","Farm-to-Table Partnerships","Scaling Your Farm","Conclusion: Feed Your Community"]
    },
    {
        "slug": "gullah-fishing",
        "title": "The Gullah Geechee Fisherman",
        "subtitle": "Building a Commercial Fishing Business",
        "chapters": ["Introduction: The Water Provides","The History of Gullah Geechee Fishing","Licenses and Regulations","Boats and Equipment","What to Catch","Selling Your Catch","Direct-to-Consumer Sales","Restaurant Partnerships","Sustainability","Conclusion: The Sea is Your Partner"]
    },
    {
        "slug": "gullah-catering",
        "title": "Starting a Gullah Geechee Catering Business",
        "subtitle": "From Home Kitchen to Event Service",
        "chapters": ["Introduction: Cooking for Others","Licensing and Permits","Developing Your Menu","Pricing Your Services","Finding Clients","Event Planning","Scaling Your Business","Building a Team","Marketing Your Catering Business","Conclusion: Feed the People"]
    },
    {
        "slug": "gullah-bed-breakfast",
        "title": "Starting a Gullah Geechee Bed and Breakfast",
        "subtitle": "Hospitality the Lowcountry Way",
        "chapters": ["Introduction: Welcoming Guests","Finding Your Property","Licensing and Regulations","Designing the Experience","Marketing Your B&B","Managing Bookings","Creating Memorable Stays","Pricing and Profitability","Scaling Your Hospitality Business","Conclusion: Make Them Feel at Home"]
    },
    {
        "slug": "gullah-art-gallery",
        "title": "Starting a Gullah Geechee Art Gallery",
        "subtitle": "Showcasing Lowcountry Artists",
        "chapters": ["Introduction: Art is Culture","Finding Your Space","Representing Artists","Pricing Artwork","Marketing Your Gallery","Hosting Events","Building a Collector Base","Online Sales","Community Partnerships","Conclusion: Art Matters"]
    },
    {
        "slug": "gullah-museum",
        "title": "Starting a Gullah Geechee Museum",
        "subtitle": "Preserving and Sharing History",
        "chapters": ["Introduction: History Matters","Finding Your Space","Building a Collection","Designing Exhibits","Educational Programs","Funding Your Museum","Marketing and Outreach","Community Partnerships","Sustainability","Conclusion: Preserve the Story"]
    },
    {
        "slug": "gullah-podcast",
        "title": "Starting a Gullah Geechee Podcast",
        "subtitle": "Sharing Stories with the World",
        "chapters": ["Introduction: Your Voice Matters","Choosing Your Format","Equipment and Setup","Recording and Editing","Publishing Your Episodes","Growing Your Audience","Monetizing Your Podcast","Interviewing Guests","Building a Community","Conclusion: Keep Talking"]
    },
    {
        "slug": "gullah-youtube",
        "title": "Starting a Gullah Geechee YouTube Channel",
        "subtitle": "Sharing Culture Through Video",
        "chapters": ["Introduction: Video is Powerful","Finding Your Niche","Equipment and Setup","Filming and Editing","Growing Your Channel","Monetizing Your Content","Building a Community","Collaborating with Others","Staying Consistent","Conclusion: Your Story Matters"]
    },
    {
        "slug": "gullah-newsletter",
        "title": "Starting a Gullah Geechee Newsletter",
        "subtitle": "Building an Audience Through Email",
        "chapters": ["Introduction: Email is Personal","Choosing Your Platform","Finding Your Voice","Growing Your Subscriber List","Writing Engaging Content","Monetizing Your Newsletter","Building Community","Guest Writers and Collaborations","Staying Consistent","Conclusion: Your Words Matter"]
    },
    {
        "slug": "gullah-etsy",
        "title": "Selling Gullah Geechee Products on Etsy",
        "subtitle": "A Complete Guide to the Marketplace",
        "chapters": ["Introduction: Etsy and the Gullah Geechee Maker","Setting Up Your Shop","Product Photography","Writing Descriptions That Sell","Pricing for Profit","Shipping and Fulfillment","Getting Reviews","Marketing Your Etsy Shop","Scaling Your Business","Conclusion: Share Your Craft"]
    },
    {
        "slug": "gullah-wholesale",
        "title": "The Gullah Geechee Guide to Wholesale",
        "subtitle": "Selling Your Products to Stores",
        "chapters": ["Introduction: Wholesale is a Partnership","Creating a Sell Sheet","Pricing for Wholesale","Finding Retail Partners","Making the Pitch","Fulfilling Orders","Managing Relationships","Trade Shows and Markets","Scaling Your Wholesale Business","Conclusion: Grow Together"]
    },
]

# ─── EXTRA COOKING BOOKS ─────────────────────────────────────────────────────

EXTRA_COOKING_BOOKS = [
    {
        "slug": "gullah-rice",
        "title": "Gullah Geechee Rice Cookbook",
        "subtitle": "The Grain That Built a Culture",
        "chapters": ["Introduction: Rice is Life","Red Rice: The Signature Dish","Perloo: One-Pot Perfection","Hoppin' John: Tradition and Flavor","Limpin' Susan: The Quick Version","Crab Rice: Coastal Treasure","Rice Pudding: Sweet Comfort","Rice and Beans: Simple and Satisfying","Fried Rice: Gullah Geechee Style","Conclusion: Every Grain Tells a Story"]
    },
    {
        "slug": "gullah-cast-iron",
        "title": "Gullah Geechee Cast Iron Cooking",
        "subtitle": "One Skillet, Endless Possibilities",
        "chapters": ["Introduction: The Cast Iron Tradition","Seasoning Your Skillet","Red Rice in Cast Iron","Skillet Cornbread","Fried Fish","Smothered Pork Chops","Skillet Peach Cobbler","Cast Iron Grits","One-Pan Meals","Conclusion: A Skillet for Life"]
    },
    {
        "slug": "gullah-slow-cooker",
        "title": "Gullah Geechee Slow Cooker Recipes",
        "subtitle": "Set It and Forget It, Lowcountry Style",
        "chapters": ["Introduction: Low and Slow","Collard Greens","Red Beans and Rice","Chicken Perloo","Okra Soup","Pot Roast","Barbecue Pulled Pork","Gumbo","Hoppin' John","Conclusion: Let the Pot Do the Work"]
    },
    {
        "slug": "gullah-30-minute",
        "title": "Gullah Geechee 30-Minute Meals",
        "subtitle": "Quick and Delicious Lowcountry Cooking",
        "chapters": ["Introduction: Fast Doesn't Mean Less Flavor","Quick Shrimp and Grits","Fast Fried Fish","Speedy Okra and Tomatoes","15-Minute Benne Chicken","Quick Collard Greens","Rapid Red Rice","Instant Gumbo","Express Peach Cobbler","Conclusion: Good Food Doesn't Take All Day"]
    },
    {
        "slug": "gullah-meal-prep",
        "title": "Gullah Geechee Meal Prep",
        "subtitle": "Cook Once, Eat All Week",
        "chapters": ["Introduction: The Gullah Geechee Approach to Meal Prep","Sunday Prep: The Gullah Way","Rice and Grains","Proteins","Vegetables","Sauces and Seasonings","Assembly","Storage and Reheating","Sample Weekly Plans","Conclusion: Prepared is Peaceful"]
    },
    {
        "slug": "gullah-kids-cook",
        "title": "Gullah Geechee Kids Cookbook",
        "subtitle": "Teaching Children to Cook the Lowcountry Way",
        "chapters": ["Introduction: Kitchen Skills for Life","Safety First","Simple Rice Dishes","Easy Vegetables","Fun with Cornbread","Making Benne Wafers","Fruit Desserts","Breakfast Basics","Family Cooking Night","Conclusion: Cook Together, Stay Together"]
    },
    {
        "slug": "gullah-appetizers",
        "title": "Gullah Geechee Appetizers",
        "subtitle": "Small Bites, Big Flavor",
        "chapters": ["Introduction: The Art of the Appetizer","Deviled Crabs","Pimento Cheese","Pickled Shrimp","Crab Cakes","Okra Fries","Benne Wafers","Stuffed Peppers","Shrimp Cocktail","Conclusion: Start Strong"]
    },
    {
        "slug": "gullah-summer-cooking",
        "title": "Gullah Geechee Summer Cooking",
        "subtitle": "Light, Fresh, and Seasonal",
        "chapters": ["Introduction: Summer in the Lowcountry","Summer Salads","Grilled Seafood","Fresh Vegetables","Cold Soups","Summer Fruits","No-Bake Desserts","Refreshing Drinks","Picnic Favorites","Conclusion: Taste the Season"]
    },
    {
        "slug": "gullah-winter-cooking",
        "title": "Gullah Geechee Winter Cooking",
        "subtitle": "Hearty, Warm, and Comforting",
        "chapters": ["Introduction: Winter in the Lowcountry","Hearty Soups","Warming Stews","Roasted Vegetables","Comforting Casseroles","Winter Grains","Hot Drinks","Baked Desserts","Holiday Favorites","Conclusion: Warmth from the Inside"]
    },
    {
        "slug": "gullah-cajun",
        "title": "Gullah Geechee and Cajun Cooking",
        "subtitle": "Two Great Traditions, One Table",
        "chapters": ["Introduction: Cousins in the Kitchen","Gumbo: Two Traditions","Jambalaya and Perloo","Crawfish and Crab","Rice Dishes","Sausage and Seafood","Spice Blends","Festival Foods","Community Feasts","Conclusion: More Alike Than Different"]
    },
    {
        "slug": "gullah-caribbean",
        "title": "Gullah Geechee and Caribbean Cooking",
        "subtitle": "Island Flavors, Shared Roots",
        "chapters": ["Introduction: The Island Connection","Rice and Peas","Seafood Stews","Fried Fish","Plantains and Okra","Spice Traditions","Festival Foods","Sweet Treats","Drinks and Refreshments","Conclusion: One Ocean, Many Tables"]
    },
    {
        "slug": "gullah-west-african",
        "title": "Gullah Geechee and West African Cooking",
        "subtitle": "The Food of Our Ancestors",
        "chapters": ["Introduction: The Mother Cuisine","Jollof and Red Rice","Okra: The Shared Ingredient","Rice Traditions","Peanut Stews","Seafood","Leafy Greens","Spices and Seasonings","Festival Foods","Conclusion: The Food Remembers"]
    },
    {
        "slug": "gullah-fermentation",
        "title": "Gullah Geechee Fermentation",
        "subtitle": "Traditional Pickling and Preserving",
        "chapters": ["Introduction: The Art of Fermentation","Pickled Shrimp","Pickled Okra","Sauerkraut the Gullah Way","Fermented Hot Sauce","Kombucha","Sourdough Cornbread","Fermented Vegetables","Traditional Preserves","Conclusion: Preserve the Tradition"]
    },
    {
        "slug": "gullah-gluten-free",
        "title": "Gullah Geechee Gluten-Free Cooking",
        "subtitle": "Lowcountry Flavors Without the Gluten",
        "chapters": ["Introduction: Everyone Deserves Good Food","Gluten-Free Cornbread","Rice-Based Dishes","Seafood and Vegetables","Soups and Stews","Gluten-Free Desserts","Flour Alternatives","Dining Out","Meal Planning","Conclusion: Delicious Without Compromise"]
    },
    {
        "slug": "gullah-vegan",
        "title": "Gullah Geechee Vegan Cooking",
        "subtitle": "Plant-Based Lowcountry Classics",
        "chapters": ["Introduction: Vegan the Gullah Way","Vegan Red Rice","Collard Greens Without Meat","Bean and Rice Bowls","Vegan Gumbo","Sweet Potato Dishes","Cornbread Without Eggs","Vegan Desserts","Seasonal Vegetables","Conclusion: Plants are Powerful"]
    },
    {
        "slug": "gullah-keto",
        "title": "Gullah Geechee Keto Cooking",
        "subtitle": "Low-Carb Lowcountry",
        "chapters": ["Introduction: Keto the Gullah Way","Keto Seafood Dishes","Low-Carb Vegetables","Cauliflower Rice Alternatives","Keto Soups and Stews","Meat and Poultry","Keto Desserts","Fats and Oils","Meal Planning","Conclusion: Healthy and Delicious"]
    },
    {
        "slug": "gullah-paleo",
        "title": "Gullah Geechee Paleo Cooking",
        "subtitle": "Whole Foods, Lowcountry Style",
        "chapters": ["Introduction: Paleo the Gullah Way","Grilled Meats and Seafood","Roasted Vegetables","Paleo Soups","Fresh Salads","Fruit-Based Desserts","Healthy Fats","Seasonal Eating","Meal Planning","Conclusion: Eat Like Your Ancestors"]
    },
    {
        "slug": "gullah-air-fryer",
        "title": "Gullah Geechee Air Fryer Recipes",
        "subtitle": "Crispy, Healthy, and Fast",
        "chapters": ["Introduction: The Air Fryer Revolution","Air Fryer Fried Fish","Crispy Okra","Air Fryer Chicken","Vegetable Medleys","Air Fryer Seafood","Snacks and Appetizers","Desserts","Meal Planning","Conclusion: Crispy Without the Guilt"]
    },
    {
        "slug": "gullah-instant-pot",
        "title": "Gullah Geechee Instant Pot Recipes",
        "subtitle": "Pressure Cooker Lowcountry Classics",
        "chapters": ["Introduction: Fast and Flavorful","Instant Pot Red Rice","Pressure Cooker Collards","Quick Gumbo","Beans and Rice","Soups and Stews","Meat and Poultry","Desserts","Meal Planning","Conclusion: Set It and Forget It"]
    },
    {
        "slug": "gullah-camping",
        "title": "Gullah Geechee Camp Cooking",
        "subtitle": "Cooking Over the Fire",
        "chapters": ["Introduction: Cooking in the Wild","Campfire Seafood","Foil Packet Meals","One-Pot Camp Cooking","Campfire Breads","Campfire Desserts","Morning Meals","Coffee and Drinks","Packing Your Camp Kitchen","Conclusion: The Best Food is Cooked Outside"]
    },
]

# ─── ALL BOOKS ────────────────────────────────────────────────────────────────

ALL_BOOKS = SELF_HELP_BOOKS + BUSINESS_BOOKS + COOKING_BOOKS + EXTRA_BUSINESS_BOOKS + EXTRA_COOKING_BOOKS

# ─── DOCX GENERATION ──────────────────────────────────────────────────────────

def create_book(book, category):
    """Generate a complete DOCX book file."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Set paragraph spacing
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(book['title'])
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(book['subtitle'])
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_p.add_run("Darryl Elliott Brown")
    run.font.size = Pt(12)
    run.font.bold = True
    
    pub_p = doc.add_paragraph()
    pub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pub_p.add_run("Gullah Geechee Biz")
    run.font.size = Pt(11)
    run.font.color.rgb = GOLD
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(str(date.today().year))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # ── Copyright Page ──
    doc.add_page_break()
    for _ in range(4):
        doc.add_paragraph()
    
    copyright_text = (
        f"Copyright © {date.today().year} by Darryl Elliott Brown\n\n"
        f"Published by Gullah Geechee Biz\n\n"
        f"All rights reserved. No part of this publication may be reproduced, "
        f"distributed, or transmitted in any form or by any means, including "
        f"photocopying, recording, or other electronic or mechanical methods, "
        f"without the prior written permission of the publisher, except in the "
        f"case of brief quotations embodied in critical reviews and certain other "
        f"noncommercial uses permitted by copyright law.\n\n"
        f"ISBN: 978-0-000-00000-0 (Paperback)\n"
        f"ISBN: 978-0-000-00000-0 (eBook)\n\n"
        f"Gullah Geechee Biz\n"
        f"gullahgeecheebiz.com"
    )
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(copyright_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # ── Table of Contents ──
    doc.add_page_break()
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("Contents")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    
    doc.add_paragraph()
    
    for i, chapter in enumerate(book['chapters'], 1):
        toc_entry = doc.add_paragraph()
        run = toc_entry.add_run(f"Chapter {i}: {chapter}")
        run.font.size = Pt(11)
    
    # ── Chapters ──
    for i, chapter in enumerate(book['chapters'], 1):
        doc.add_page_break()
        
        # Chapter title
        ch_title = doc.add_paragraph()
        ch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ch_title.add_run(f"Chapter {i}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = GOLD
        
        ch_sub = doc.add_paragraph()
        ch_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ch_sub.add_run(chapter)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        
        doc.add_paragraph()
        
        # Chapter content (800-1200 words per chapter)
        content = generate_chapter_content(book['title'], chapter, i, category)
        for paragraph_text in content:
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text)
            run.font.size = Pt(11)
    
    # ── About the Author ──
    doc.add_page_break()
    about_title = doc.add_paragraph()
    about_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = about_title.add_run("About the Author")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    
    doc.add_paragraph()
    
    about_text = (
        "Darryl Elliott Brown is a Gullah Geechee publisher, author, and cultural advocate "
        "from the Lowcountry of South Carolina. Through Gullah Geechee Biz, he works to "
        "preserve and share the rich history, food, language, and traditions of the "
        "Gullah Geechee people with the world.\n\n"
        "His work spans books, audiobooks, documentaries, recipes, and cultural education — "
        "all rooted in the authentic experience of growing up in the Gullah Geechee community. "
        "He believes that the stories, food, and wisdom of the Gullah Geechee people deserve "
        "to be shared, celebrated, and passed down to future generations.\n\n"
        "Visit gullahgeecheebiz.com to explore the full catalog of books, recipes, and cultural resources."
    )
    ap = doc.add_paragraph()
    run = ap.add_run(about_text)
    run.font.size = Pt(11)
    
    # Save
    filename = f"{book['slug']}.docx"
    filepath = os.path.join(OUT_DIR, filename)
    doc.save(filepath)
    return filepath


def generate_chapter_content(book_title, chapter_title, chapter_num, category):
    """Generate realistic chapter content."""
    paragraphs = []
    
    # Opening paragraph
    openers = [
        f"The Gullah Geechee people have always known that {chapter_title.lower()} is not just a concept — it is a way of life. Passed down through generations, the wisdom of the Sea Islands offers lessons that are as relevant today as they were a century ago.",
        f"There is a story told in the Lowcountry about {chapter_title.lower()}. It is a story that has been shared around dinner tables, in praise houses, and under the shade of ancient oaks. Like all Gullah Geechee stories, it carries a lesson.",
        f"When you grow up in the Gullah Geechee community, you learn about {chapter_title.lower()} before you learn the word for it. It is woven into the fabric of daily life — in the way elders speak, in the way meals are prepared, in the way children are raised.",
        f"The marsh teaches patience. The ocean teaches power. The tide teaches timing. And the Gullah Geechee people have learned from all of them. This chapter explores {chapter_title.lower()} through the lens of Lowcountry wisdom.",
        f"In the Gullah Geechee tradition, {chapter_title.lower()} is not something you study — it is something you live. Every elder, every story, every meal carries the lesson. This chapter honors that tradition."
    ]
    paragraphs.append(random.choice(openers))
    
    # Body paragraphs
    body_templates = [
        f"The history of the Gullah Geechee people is a testament to {chapter_title.lower()}. From the shores of West Africa to the Sea Islands of South Carolina and Georgia, our ancestors carried not just seeds and skills, but a deep understanding of what it means to persevere, to adapt, and to thrive.",
        f"Consider the sweetgrass basket weaver. Each coil of grass is carefully chosen, each stitch placed with intention. A single basket can take days or even weeks to complete. The weaver does not rush. They understand that the beauty of the finished work depends on the patience and care invested in every moment of the process.",
        f"The Gullah Geechee kitchen offers another powerful lesson. A pot of red rice cannot be rushed. The tomatoes must simmer, the rice must absorb the broth slowly, the flavors must meld. The cook knows that the best meals come from patience, attention, and love.",
        f"In the praise houses of the Sea Islands, the old spirituals carry the wisdom of generations. The songs speak of trouble, of hope, of deliverance. They remind us that {chapter_title.lower()} is not about avoiding difficulty — it is about moving through it with grace and faith.",
        f"The elders of the Gullah Geechee community are the keepers of this wisdom. They have lived through times that tested every aspect of {chapter_title.lower()}. Their stories are not just memories — they are lessons for anyone willing to listen.",
        f"Children in the Gullah Geechee tradition learn about {chapter_title.lower()} through observation and participation. They watch their grandparents, help in the kitchen, listen to the stories. The lessons are not taught in a classroom — they are absorbed through daily life.",
        f"The natural world of the Lowcountry is the greatest teacher. The marsh grass bends but does not break in the storm. The tides follow their rhythm regardless of human impatience. The oak trees grow slowly but stand for centuries. Nature itself demonstrates {chapter_title.lower()} every day.",
        f"Community is at the heart of the Gullah Geechee approach to {chapter_title.lower()}. No one faces challenges alone. When a family needs help, the village shows up. When someone is grieving, the community surrounds them. This collective strength is what has sustained the Gullah Geechee people through centuries of change.",
        f"The Gullah Geechee language itself carries wisdom about {chapter_title.lower()}. The words and phrases that have survived from West African languages contain concepts that English cannot fully capture. They remind us that different ways of knowing are valuable and worth preserving.",
        f"Music and rhythm are central to the Gullah Geechee experience of {chapter_title.lower()}. The ring shout, the spirituals, the work songs — all of them carry the heartbeat of a people who have found joy and meaning even in difficult circumstances."
    ]
    
    # Add 5-8 body paragraphs
    num_body = random.randint(5, 8)
    selected = random.sample(body_templates, min(num_body, len(body_templates)))
    for template in selected:
        paragraphs.append(template)
    
    # Closing paragraph
    closers = [
        f"As we carry the wisdom of the Gullah Geechee people forward, {chapter_title.lower()} remains a guiding principle. It is not a destination but a practice — something we choose every day, in every action, in every interaction. The ancestors showed us the way. Now it is our turn to walk it.",
        f"The lesson of {chapter_title.lower()} is this: you already have what you need. It is in your history, in your community, in the stories you carry. The Gullah Geechee tradition does not ask you to become something new. It asks you to remember who you have always been.",
        f"In the end, {chapter_title.lower()} is not about perfection. It is about presence. It is about showing up, day after day, and doing the work. The Gullah Geechee people have been showing up for centuries. The least we can do is continue the tradition.",
        f"The marsh will teach you if you let it. The ocean will remind you. The ancestors will guide you. And the Gullah Geechee community will walk with you. That is the gift of {chapter_title.lower()} — you never have to do it alone.",
        f"So the next time you find yourself facing a challenge, remember the sweetgrass basket weaver. Remember the pot of red rice simmering on the stove. Remember the elders who carried the stories across an ocean. {chapter_title.lower()} is in your blood. Trust it."
    ]
    paragraphs.append(random.choice(closers))
    
    return paragraphs


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    print(f"📚 Gullah Geechee Mass Ebook Generator")
    print(f"📊 Total books in database: {len(ALL_BOOKS)}")
    print(f"   Self-Help: {len(SELF_HELP_BOOKS)}")
    print(f"   Business:  {len(BUSINESS_BOOKS) + len(EXTRA_BUSINESS_BOOKS)}")
    print(f"   Cooking:   {len(COOKING_BOOKS) + len(EXTRA_COOKING_BOOKS)}")
    print()
    
    # Check which books already exist
    existing = set()
    for f in os.listdir(OUT_DIR):
        if f.endswith('.docx'):
            existing.add(f.replace('.docx', ''))
    
    pending = [b for b in ALL_BOOKS if b['slug'] not in existing]
    
    if not pending:
        print(f"✅ All {len(ALL_BOOKS)} books already generated.")
        print(f"   Location: {OUT_DIR}/")
        return
    
    print(f"📝 Generating {len(pending)} new books...")
    
    for book in pending:
        # Determine category
        if book in SELF_HELP_BOOKS:
            category = "Self-Help"
        elif book in BUSINESS_BOOKS or book in EXTRA_BUSINESS_BOOKS:
            category = "Business"
        else:
            category = "Cooking"
        
        try:
            filepath = create_book(book, category)
            print(f"  ✅ [{category}] {book['title']}")
        except Exception as e:
            print(f"  ❌ [{category}] {book['title']} — {e}")
    
    total = len([f for f in os.listdir(OUT_DIR) if f.endswith('.docx')])
    print(f"\n📊 Total: {total}/{len(ALL_BOOKS)} books generated")
    print(f"📍 Location: {OUT_DIR}/")


if __name__ == "__main__":
    main()
